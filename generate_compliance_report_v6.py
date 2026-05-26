"""
K Recruit — HR Resume Optimizer
DATA PRIVACY & COMPLIANCE ASSESSMENT REPORT  VERSION 6
Document Ref: K-RECRUIT-COMP-2026-006

COMPLETE document — fully supersedes v1–v5 (2026-001 through 2026-005).

New in v6 (May 2026):
  - Section 2 updated: Dedicated office laptop named as the production server
  - Section 19: Server Laptop Deployment Guide (hardware, PostgreSQL, static IP, power)
  - Section 20: Operational Procedures (backup, log rotation, firewall, boot)
  - New findings added: APP_ENV not set, data isolation missing, backup missing,
                        firewall gap, log rotation gap
  - Backlog section added: password policy, account lifecycle, BitLocker
  - Section 14 checklist updated with server laptop specific items
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

NAVY    = RGBColor(0x0F, 0x17, 0x2A)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
MUTED   = RGBColor(0x64, 0x74, 0x8B)
TEAL    = RGBColor(0x0D, 0x6E, 0x6E)
RED     = RGBColor(0xC0, 0x10, 0x20)
GREEN   = RGBColor(0x16, 0x65, 0x34)
AMBER   = RGBColor(0x92, 0x40, 0x0E)
GRAY_L  = RGBColor(0x94, 0xA3, 0xB8)
INDIGO_L= RGBColor(0x99, 0xA3, 0xFF)
INDIGO  = RGBColor(0x44, 0x38, 0xCA)

def shade_cell(cell, hex6):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6.lstrip("#")); tcPr.append(shd)

def cell_pad(cell, twips=140):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ("top","left","bottom","right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(twips)); m.set(qn("w:type"), "dxa"); tcMar.append(m)
    tcPr.append(tcMar)

def bottom_border(para, color="0F172A"):
    pPr = para._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "2"); bot.set(qn("w:color"), color)
    pBdr.append(bot); pPr.append(pBdr)

def left_bar(cell, color="6366F1"):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBorders")
    for edge in ("top","bottom","right"):
        t = OxmlElement(f"w:{edge}"); t.set(qn("w:val"), "none"); tcBdr.append(t)
    lft = OxmlElement("w:left")
    lft.set(qn("w:val"), "single"); lft.set(qn("w:sz"), "18")
    lft.set(qn("w:space"), "0"); lft.set(qn("w:color"), color)
    tcBdr.append(lft); tcPr.append(tcBdr)

def rn(para, text, bold=False, italic=False, sz=10, color=None):
    r = para.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(sz)
    if color: r.font.color.rgb = color
    return r

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
    bottom_border(p); return p

def h2(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11.5)
    r.font.color.rgb = color or NAVY; return p

def h3(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(10.5)
    r.font.color.rgb = color or MUTED; return p

def body(text, sz=10, color=None, after=5, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(sz); r.bold = bold; r.italic = italic
    r.font.color.rgb = color or NAVY; return p

def bul(text, sz=10, indent=0.35):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text); r.font.size = Pt(sz); r.font.color.rgb = NAVY; return p

def sp(pts=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts)

def callout(lines, bg="EEF2FF", border="6366F1"):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0,0); shade_cell(cell, bg); left_bar(cell, border); cell_pad(cell, 200)
    cell.paragraphs[0].clear()
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i==0 else cell.add_paragraph()
        r = p.add_run(line); r.font.size = Pt(9.5); r.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    sp(6)

def tbl(headers, rows, col_widths=None, hdr_bg="0F172A", alt="F8FAFC"):
    nc = len(headers)
    t = doc.add_table(rows=1+len(rows), cols=nc)
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hr = t.rows[0]
    for j,h in enumerate(headers):
        c = hr.cells[j]; shade_cell(c, hdr_bg); cell_pad(c, 120)
        p = c.paragraphs[0]; r = p.add_run(h)
        r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    for i, row in enumerate(rows):
        bg = alt if i%2==0 else "FFFFFF"; tr = t.rows[i+1]
        for j, cv in enumerate(row):
            c = tr.cells[j]; shade_cell(c, bg); cell_pad(c, 120)
            p = c.paragraphs[0]
            if isinstance(cv, tuple):
                txt, bold, clr = cv
                r = p.add_run(str(txt)); r.bold = bold
                r.font.size = Pt(9); r.font.color.rgb = clr
            else:
                r = p.add_run(str(cv)); r.font.size = Pt(9); r.font.color.rgb = NAVY
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    if col_widths:
        for j,w in enumerate(col_widths):
            for row in t.rows: row.cells[j].width = Inches(w)
    sp(6)

def OK(t):    return (t, True,  GREEN)
def BAD(t):   return (t, True,  RED)
def WRN(t):   return (t, True,  AMBER)
def DONE(t):  return (t, True,  GREEN)
def BKLG(t):  return (t, True,  INDIGO)

# =============================================================================
#  COVER PAGE
# =============================================================================
ct = doc.add_table(rows=1, cols=1)
cc = ct.cell(0,0); shade_cell(cc, "0F172A"); cell_pad(cc, 380)
cc.paragraphs[0].clear()

def cl(text, sz, clr, bold=False, after=6):
    p = cc.add_paragraph(); r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(sz); r.font.color.rgb = clr
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(after)

cl("K Recruit",                                                          24, WHITE,    bold=True,  after=2)
cl("HR Resume Optimizer — Kforce Internal Tool",                         14, GRAY_L,   bold=False, after=16)
cl("DATA PRIVACY & COMPLIANCE ASSESSMENT REPORT",                        12, INDIGO_L, bold=True,  after=4)
cl("VERSION 6 — SERVER LAPTOP & OPERATIONAL GAPS EDITION  (Supersedes v1–v5)", 9, GRAY_L, bold=True, after=16)

for lbl, val in [
    ("Document Ref:",    "K-RECRUIT-COMP-2026-006"),
    ("Version:",         "6.0 — Final  |  Supersedes v1–v5 (2026-001 through 2026-005)"),
    ("Prepared for:",    "Kforce US Compliance / IT Governance Team"),
    ("Prepared by:",     "Application Compliance Review — Technical Architecture"),
    ("Scope:",           "Dedicated server laptop deployment + operational gaps identified and classified"),
    ("Review Date:",     "May 2026"),
    ("Classification:",  "Internal Confidential"),
]:
    p = cc.add_paragraph()
    r1 = p.add_run(f"{lbl:20s}"); r1.bold = True; r1.font.size = Pt(9); r1.font.color.rgb = GRAY_L
    r2 = p.add_run(val);          r2.font.size = Pt(9);                  r2.font.color.rgb = WHITE
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(3)

cc.paragraphs[0]._element.getparent().remove(cc.paragraphs[0]._element)
sp(12)

# =============================================================================
#  SECTION 1 — EXECUTIVE SUMMARY
# =============================================================================
h1("1.  Executive Summary")

body(
    "This report is the definitive compliance and data privacy assessment for K Recruit, "
    "Kforce's internal HR resume optimisation tool. It supersedes all previous versions "
    "(v1–v5: 2026-001 through 2026-005). "
    "A complete line-by-line review was conducted across all backend services, routers, "
    "data models, configuration files, and frontend code."
)

callout([
    "KEY FINDING: K Recruit is architecturally capable of operating with ZERO DATA EGRESS",
    "when configured with Microsoft Azure OpenAI (Kforce's own Azure tenant).",
    "",
    "The current development configuration uses third-party cloud APIs (Groq, Gemini, NVIDIA NIM)",
    "which transmit candidate PII outside the organisation. This is a DEVELOPMENT-ONLY",
    "configuration and must be switched to Azure OpenAI before any production use.",
    "",
    "The single most important action: IT provisions Azure OpenAI in Kforce's Azure tenant.",
    "This is a Settings change in the app — no code deployment required.",
], bg="EEF2FF", border="6366F1")

body("Versions 4 and 5 resolved all critical authentication and security hardening findings:")
for item in [
    "All API routes require JWT authentication — including /bot/chat (fixed in v5)",
    "JWT_SECRET loaded from environment variable — not hardcoded in source code",
    "Passwords stored as bcrypt hashes in users table — no plain text credentials",
    "JWT tokens expire after 24 hours",
    "Brute-force login protection: 5 failed attempts per IP per 60 seconds → HTTP 429",
    "Audit logging middleware: all HTTP requests logged to audit.log",
    "SRI integrity hashes on all 3 CDN script tags",
    "PII stripped from AI prompts (email, phone, LinkedIn) before LLM call",
    "Full multi-user RBAC: super_admin / admin / recruiter with permission sets",
]:
    bul(item)

sp(4)
body(
    "Version 6 — Server Laptop & Operational Gaps Edition — confirms the dedicated office laptop "
    "as the production server and documents five new compliance findings identified "
    "during the 100-recruiter deployment review.",
    bold=True
)
for item in [
    "New finding: APP_ENV not set — FastAPI /docs and /redoc remain accessible in default mode",
    "New finding: Candidate data isolation missing — all 100 recruiters can see all candidates",
    "New finding: No backup strategy — single point of failure on the server laptop",
    "New finding: Windows Firewall gap — port 8000 accessible directly, bypassing nginx/SSL",
    "New finding: No audit log rotation — audit.log grows indefinitely",
    "Backlog documented: password minimum length, account lifecycle policy, BitLocker",
]:
    bul(item)

# =============================================================================
#  SECTION 2 — APPLICATION OVERVIEW
# =============================================================================
h1("2.  Application Overview")

tbl(
    ["Item", "Detail"],
    [
        ("Application Name",    "K Recruit — HR Resume Optimizer"),
        ("Previous Name",       "CMS TalentAI (renamed to K Recruit in v3 of this report)"),
        ("Purpose",             "Automate resume screening, ATS gap analysis, JD alignment, and soft-gap detection for HR recruiters"),
        ("Users",               "Internal HR recruiters — Kforce India office (pilot: ~20, full rollout: ~100)"),
        ("Production Server",   "Dedicated office laptop — permanently powered, sleep disabled, Kforce India office LAN"),
        ("Server Spec (min)",   "16 GB RAM, 512 GB SSD — AI processing offloaded to Azure, server handles web + DB only"),
        ("Network",             "Fixed static LAN IP (e.g. 192.168.1.50) — assigned by IT in router/adapter settings"),
        ("Backend",             "Python 3  ·  FastAPI web framework  ·  Uvicorn ASGI server"),
        ("Frontend",            "Single-page HTML app — React 18.2 + Babel (CDN with SRI hashes)"),
        ("Database",            "SQLite (development / up to ~30 users)  ·  PostgreSQL (production — 100 users)"),
        ("AI Provider",         "Development: Groq / Gemini / NVIDIA NIM  ·  Production: Azure OpenAI ONLY"),
        ("Semantic Model",      "sentence-transformers all-MiniLM-L6-v2 (~22 MB, runs fully locally on server laptop)"),
        ("Authentication",      "Email + bcrypt password  ·  JWT token (24h expiry)  ·  Role-based access control"),
        ("User Roles",          "super_admin / admin / recruiter — full RBAC with per-role permission sets"),
        ("File Storage",        "Uploaded PDFs stored in backend/uploads/ on the server laptop — local disk only"),
        ("Vector Database",     "None — similarity scoring computed in-memory only; no embeddings persisted"),
    ],
    col_widths=[1.8, 4.7],
)

# =============================================================================
#  SECTION 3 — DATA INVENTORY
# =============================================================================
h1("3.  Data Inventory — What Personal Data K Recruit Handles")

YES_DEV = ("YES — to AI provider (dev mode only)", True, RED)
AZ_ONLY = ("To Kforce Azure tenant only (production)", True, AMBER)
NO_GRN  = ("No — local storage only", False, GREEN)

tbl(
    ["Data Element", "Source", "Where Stored", "Dev: Transmitted?", "Production: Transmitted?"],
    [
        ("Candidate Full Name",   "PDF Resume",     "Local SQLite/PG DB",         YES_DEV, AZ_ONLY),
        ("Email Address",         "PDF Resume",     "Local SQLite/PG DB",         YES_DEV, AZ_ONLY),
        ("Mobile Number",         "PDF Resume",     "Local SQLite/PG DB",         YES_DEV, AZ_ONLY),
        ("Location / City",       "PDF Resume",     "Local SQLite/PG DB",         YES_DEV, AZ_ONLY),
        ("LinkedIn URL",          "PDF Resume",     "Local SQLite/PG DB",         YES_DEV, AZ_ONLY),
        ("Full Resume Text",      "PDF Resume",     "SQLite/PG DB + /uploads",    YES_DEV, AZ_ONLY),
        ("Job Description Text",  "HR team input",  "Local SQLite/PG DB",         YES_DEV, AZ_ONLY),
        ("AI Analysis Results",   "Generated",      "Local SQLite/PG DB",         NO_GRN,  NO_GRN),
        ("Recruiter Name",        "HR team input",  "Local SQLite/PG DB",         NO_GRN,  NO_GRN),
        ("Login Credentials",     "users table",    "Local DB — bcrypt hash only",NO_GRN,  NO_GRN),
        ("PDF Files",             "Upload",         "/uploads on server laptop",  NO_GRN,  NO_GRN),
    ],
    col_widths=[1.4, 1.1, 1.5, 1.6, 1.9],
)

callout([
    "DATA ISOLATION GAP (v6 NEW FINDING):",
    "Currently the candidates table has no uploaded_by field linking to the users table.",
    "This means all 100 recruiters can see all candidates uploaded by any recruiter.",
    "For a shared office deployment this is a known gap — see Section 11 Finding C-07.",
    "Fix: add uploaded_by FK to candidates table; filter results by user role in queries.",
], bg="FEF2F2", border="EF4444")

# =============================================================================
#  SECTION 4 — DATA FLOW DIAGRAMS
# =============================================================================
h1("4.  Data Flow Diagrams")

h2("4.1  Current Development State (NOT Production-Ready for PII)", color=RED)
callout([
    "CURRENT FLOW — DEVELOPMENT ONLY:",
    "",
    "  [HR Recruiter Browser]",
    "         │",
    "         ▼   HTTP port 8000  (UNENCRYPTED — no SSL)",
    "  [Server Laptop — FastAPI Backend]",
    "         │",
    "         ├── PDF Upload ──────► /uploads folder  (LOCAL DISK)",
    "         ├── Database  ──────► cms.db SQLite     (LOCAL DISK)",
    "         │",
    "         └── AI Analysis ────► DATA LEAVES PREMISES ──►",
    "               ├── Groq Cloud    (api.groq.com — San Jose, CA, USA)",
    "               ├── Google Gemini (generativelanguage.googleapis.com)",
    "               └── NVIDIA NIM    (integrate.api.nvidia.com)",
], bg="FEF2F2", border="EF4444")

h2("4.2  Production-Ready State — Azure OpenAI + SSL + Server Laptop", color=GREEN)
callout([
    "PRODUCTION FLOW — Dedicated Office Laptop + nginx SSL + JWT auth:",
    "",
    "  ┌──────────────────────────────────────────────────────────────────┐",
    "  │                     KFORCE INDIA OFFICE LAN                      │",
    "  │                                                                  │",
    "  │  [Recruiter PC]  ─── HTTPS/443 (SSL) ──► [nginx on laptop]      │",
    "  │  (any of 100)                                     │              │",
    "  │                                      HTTP/8000 (localhost only)  │",
    "  │                                                   │              │",
    "  │                                    ┌──────────────▼───────────┐  │",
    "  │  PDF uploaded ────────────────────►│  FastAPI — Server Laptop  │ │",
    "  │  JWT token on every request        │  PostgreSQL DB            │ │",
    "  │  Static IP: 192.168.1.50           │  /uploads (local disk)    │ │",
    "  │  Port 8000 blocked at firewall     └──────────┬────────────────┘ │",
    "  └─────────────────────────────────────────────  │  ────────────────┘",
    "                                                   │ Resume + JD text",
    "                                                   ▼  HTTPS/443 only",
    "                               ┌──────────────────────────────────────┐",
    "                               │   KFORCE AZURE TENANT                │",
    "                               │   Azure OpenAI — GPT-4o              │",
    "                               │   Microsoft DPA applies              │",
    "                               └──────────────────────────────────────┘",
], bg="F0FDF4", border="16A34A")

# =============================================================================
#  SECTION 5 — HUGGINGFACE MODEL
# =============================================================================
h1("5.  HuggingFace Model Download — Audit")

tbl(
    ["Detail", "Value"],
    [
        ("Model Name",           "all-MiniLM-L6-v2"),
        ("License",              "Apache 2.0 — free for all commercial use"),
        ("File size",            "~22 MB"),
        ("Downloaded when",      "ONCE — on first server startup after pip install"),
        ("After first download", "Runs 100% locally — zero further network calls to HuggingFace"),
        ("Production action",    "Pre-download during setup; set HF_HUB_OFFLINE=1 in .env on server laptop"),
        ("User data sent",       "NONE — only the model weights file is downloaded"),
    ],
    col_widths=[2.2, 4.3],
)

callout([
    "NOTE: HF_HUB_OFFLINE concern is resolved automatically when Azure OpenAI is configured.",
    "Once the server laptop is in production with Azure, the semantic model is pre-downloaded",
    "during setup and HF_HUB_OFFLINE=1 is set in .env — no further HuggingFace contact.",
    "This is part of the server laptop setup checklist (Section 19).",
], bg="F0FDF4", border="16A34A")

h2("5.1  Alternatives if Compliance Does Not Approve HuggingFace")
tbl(
    ["Option", "HuggingFace?", "Quality", "Effort", "Best When"],
    [
        ("TF-IDF fallback (built-in)", OK("None"), WRN("Lower"), OK("Remove 1 line"), "Fastest, zero contact"),
        ("Internal file share",        OK("None"), OK("Identical"), WRN("Low"), "Kforce has file server"),
        ("Bundle with app",            OK("None"), OK("Identical"), WRN("Low"), "Clean enterprise deploy"),
        ("Azure embeddings",           OK("None"), OK("Better"),    WRN("~20 lines"), "Already using Azure"),
    ],
    col_widths=[1.8, 1.0, 1.0, 1.0, 2.7],
)

# =============================================================================
#  SECTION 6 — LIBRARY NETWORK AUDIT
# =============================================================================
h1("6.  Complete Library-by-Library Network Audit")

h2("6.1  Backend Python Libraries")
tbl(
    ["Library", "Version", "Purpose", "Network Calls?", "User Data?", "Verdict"],
    [
        ("fastapi",              "≥0.110", "Web API framework",          OK("None"), OK("No"), OK("SAFE")),
        ("uvicorn",              "≥0.29",  "ASGI web server",            OK("None"), OK("No"), OK("SAFE")),
        ("sqlalchemy",           "≥2.0",   "Database ORM",               OK("None"), OK("No"), OK("SAFE")),
        ("psycopg2-binary",      "≥2.9",   "PostgreSQL driver",          OK("None"), OK("No"), OK("SAFE")),
        ("pymupdf (fitz)",       "≥1.24",  "PDF text extraction",        OK("None"), OK("No"), OK("SAFE")),
        ("python-docx",          "≥1.1",   "Word document generation",   OK("None"), OK("No"), OK("SAFE")),
        ("requests",             "≥2.31",  "AI provider HTTP calls",     WRN("Azure OpenAI only"), WRN("Resume+JD"), WRN("APPROVED")),
        ("python-dotenv",        "≥1.0",   "Reads .env config file",     OK("None"), OK("No"), OK("SAFE")),
        ("python-jose",          "≥3.3",   "JWT token signing",          OK("None"), OK("No"), OK("SAFE")),
        ("passlib[bcrypt]",      "≥1.7",   "Password hashing",           OK("None"), OK("No"), OK("SAFE")),
        ("sentence-transformers","≥2.7",   "Semantic scoring",           WRN("One-time model DL"), OK("No"), WRN("MITIGATED")),
        ("scikit-learn",         "≥1.3",   "TF-IDF fallback scoring",    OK("None"), OK("No"), OK("SAFE")),
    ],
    col_widths=[1.5, 0.6, 1.7, 1.4, 1.0, 0.85],
)

h2("6.2  Frontend JavaScript Dependencies")
tbl(
    ["Library", "CDN", "Version", "SRI Hash?", "User Data Sent?", "Verdict"],
    [
        ("React",    "cdnjs.cloudflare.com", "18.2.0", OK("sha384 — verified"), OK("No"), OK("SAFE")),
        ("ReactDOM", "cdnjs.cloudflare.com", "18.2.0", OK("sha384 — verified"), OK("No"), OK("SAFE")),
        ("Babel",    "cdnjs.cloudflare.com", "7.23.2", OK("sha384 — verified"), OK("No"), OK("SAFE")),
    ],
    col_widths=[0.9, 2.0, 0.8, 1.5, 1.2, 0.8],
)

# =============================================================================
#  SECTION 7 — AZURE OPENAI
# =============================================================================
h1("7.  Azure OpenAI — Subscription, Setup, and Cost Guide")

h2("7.1  Do We Need a New Azure Subscription?")
callout([
    "NO — Kforce already has an Azure subscription (used for Microsoft 365 / Copilot).",
    "Azure OpenAI is an add-on SERVICE within the existing Kforce Azure tenant.",
    "IT needs to ENABLE it — not create a new subscription.",
], bg="EEF2FF", border="6366F1")

h2("7.2  Step-by-Step Setup for IT Team")
tbl(
    ["Step", "Action", "Where in Azure Portal", "Who"],
    [
        ("1", "Enable Azure OpenAI for the subscription",
         "Subscriptions → Resource Providers → Microsoft.CognitiveServices → Register", "IT / Cloud Admin"),
        ("2", "Create Resource Group: 'kforce-k-recruit'",
         "portal.azure.com → Resource Groups → + Create", "IT / Cloud Admin"),
        ("3", "Create Azure OpenAI resource",
         "Create Resource → search 'Azure OpenAI' → Create", "IT / Cloud Admin"),
        ("4", "Choose region: Central India (DPDP compliance)",
         "Selected during resource creation — cannot change after", "IT + Compliance"),
        ("5", "Deploy GPT-4o model",
         "Azure OpenAI resource → Model Deployments → + Deploy → GPT-4o", "IT / Cloud Admin"),
        ("6", "Copy Endpoint URL and API Key",
         "Azure OpenAI resource → Keys and Endpoint", "IT / Cloud Admin"),
        ("7", "Enter in K Recruit Settings",
         "K Recruit Settings page → Azure OpenAI → endpoint + key → Save", "Dev / HR Admin"),
    ],
    col_widths=[0.35, 2.2, 2.7, 1.1],
)

h2("7.3  Cost Estimate")
tbl(
    ["Activity", "Tokens / Request", "Volume / Month", "Estimated Cost"],
    [
        ("Resume gap analysis",     "~3,000 tokens", "200 resumes",  "~$1.80"),
        ("JD-aligned resume gen",   "~5,000 tokens", "100 resumes",  "~$1.50"),
        ("Bot assistant chat",      "~500 tokens",   "200 messages", "~$0.30"),
        ("TOTAL — 20 recruiters",   "",              "",             OK("~$3–$10 / month")),
        ("TOTAL — 100 recruiters",  "",              "",             WRN("~$30–$80 / month")),
    ],
    col_widths=[2.0, 1.8, 1.5, 2.2],
)

# =============================================================================
#  SECTION 8 — AI PROVIDER MATRIX
# =============================================================================
h1("8.  AI Provider Compliance Matrix")

tbl(
    ["Provider", "Dev Use?", "Production Use?", "Data Leaves Premises?", "Enterprise DPA?", "Reason"],
    [
        ("Azure OpenAI",   OK("✓ Allowed"), OK("✓ APPROVED"),    ("Kforce Azure tenant only", False, TEAL), OK("Yes — Microsoft DPA"), "Same boundary as M365 Copilot"),
        ("Ollama (Local)", OK("✓ Allowed"), WRN("Optional"),     OK("No — 100% local"),                    OK("N/A — fully local"),   "Zero egress"),
        ("Groq Cloud",     WRN("Dev only"), BAD("✗ BLOCKED"),    BAD("Yes — US commercial"),               BAD("Limited ToS"),        "Third-party US server — PII outside org"),
        ("Google Gemini",  WRN("Dev only"), BAD("✗ BLOCKED"),    BAD("Yes — Google servers"),              BAD("May train on data"),  "Insufficient data retention policy"),
        ("NVIDIA NIM",     WRN("Dev only"), BAD("✗ BLOCKED"),    BAD("Yes — NVIDIA cloud"),                BAD("Research-use policy"),"Policy insufficient for enterprise PII"),
    ],
    col_widths=[1.3, 0.9, 1.0, 1.5, 1.3, 1.5],
)

# =============================================================================
#  SECTION 9 — SSL / HTTPS
# =============================================================================
h1("9.  SSL / HTTPS Certificate for K Recruit — Full Guide")

body(
    "The server laptop is on the Kforce India office LAN. All 100 recruiters connect over "
    "that shared network. Without SSL, passwords, JWT tokens, and resume data travel as plain "
    "text — readable by any device on the same network."
)

h2("9.1  Recommended Option — Self-Signed Certificate + GPO Push")
callout([
    "ZERO COST. ONE-TIME SETUP. NO BROWSER WARNINGS.",
    "",
    "Step 1 — Generate certificate on the server laptop:",
    '  openssl req -x509 -newkey rsa:4096 -keyout key.pem -out cert.pem -days 3650 -nodes \\',
    '    -subj "/C=IN/ST=Maharashtra/L=Pune/O=Kforce/CN=192.168.1.50"',
    "",
    "Step 2 — IT exports cert.pem as .cer and imports into Group Policy:",
    "  GPMC → Computer Config → Windows Settings → Security Settings",
    "  → Public Key Policies → Trusted Root Certification Authorities → Import",
    "  Link GPO to the OU containing all recruiter PCs → gpupdate /force",
    "",
    "Result: All domain-joined recruiter PCs trust the certificate.",
    "        Recruiters see a green padlock — zero browser warnings.",
], bg="F0FDF4", border="16A34A")

h2("9.2  nginx Configuration")
callout([
    "  server {",
    "      listen 443 ssl;",
    "      server_name 192.168.1.50;   # server laptop static IP",
    "",
    "      ssl_certificate     /path/to/cert.pem;",
    "      ssl_certificate_key /path/to/key.pem;",
    "      ssl_protocols       TLSv1.2 TLSv1.3;",
    "",
    "      location / {",
    "          proxy_pass       http://127.0.0.1:8000;",
    "          proxy_set_header Host $host;",
    "      }",
    "  }",
    "  server {",
    "      listen 80;",
    "      server_name 192.168.1.50;",
    "      return 301 https://$host$request_uri;",
    "  }",
], bg="F1F5F9", border="94A3B8")

h2("9.3  One Code Change Required After SSL")
callout([
    "main.py — CORS allow_origins:",
    "  BEFORE:  allow_origins=['*']",
    "  AFTER:   allow_origins=['https://192.168.1.50']",
    "",
    "  Also update API_URL in index.html / preview.html to 'https://192.168.1.50'",
], bg="EEF2FF", border="6366F1")

# =============================================================================
#  SECTION 10 — SQLITE AND POSTGRESQL
# =============================================================================
h1("10.  SQLite and PostgreSQL — License, Subscription & Certificate Facts")

callout([
    "SQLite: PUBLIC DOMAIN — no license, no fee, no subscription. Suitable up to ~30 users.",
    "PostgreSQL: Free (BSD-style license). Required for 100-user production deployment.",
    "HTTPS Certificate: Required only for nginx — NOT for the database.",
    "Migration: Change DATABASE_URL in .env only — no code changes needed.",
], bg="F0FDF4", border="16A34A")

tbl(
    ["Question", "SQLite", "PostgreSQL"],
    [
        ("License / fee?",        OK("None — Public Domain"),          OK("None — free forever")),
        ("Subscription?",         OK("None — built into Python"),      OK("None — self-hosted")),
        ("Certificate needed?",   OK("No — local file"),               OK("No — nginx handles SSL")),
        ("Max users (guideline)", WRN("~30 concurrent users"),         OK("100+ users — recommended")),
        ("Encryption at rest?",   BKLG("Backlog — BitLocker on laptop"), BKLG("Backlog — pg_tde or BitLocker")),
        ("Production ready?",     WRN("Pilot only"),                   OK("Yes — required for 100 users")),
    ],
    col_widths=[2.0, 2.3, 2.2],
)

# =============================================================================
#  SECTION 11 — SECURITY FINDINGS (UPDATED V6)
# =============================================================================
h1("11.  Security Findings — Updated v6")

h2("11.1  Resolved Findings (v4 + v5)", color=GREEN)
tbl(
    ["#", "Finding", "Resolved In", "How"],
    [
        ("C-04", "Plain-text login password",           "v5", DONE("bcrypt hash in users table")),
        ("C-06", "Missing auth on API routes",          "v4+v5", DONE("get_current_user on all routes incl. /bot/chat")),
        ("H-01", "No audit log",                        "v4", DONE("Middleware logs all requests to audit.log")),
        ("H-02", "No session token expiry",             "v5", DONE("JWT expires after 24 hours")),
        ("V4-F2","Hardcoded JWT secret",                "v4", DONE("JWT_SECRET env var — server refuses to start without it")),
        ("V4-F3","/docs and /redoc in production",      "v4", DONE("Hidden when APP_ENV=production  (see C-07 below)")),
        ("V4-F6","CDN scripts without integrity check", "v4", DONE("SRI sha384 hashes on all 3 script tags")),
        ("V4-F7","PII in AI prompts",                   "v4", DONE("Email, phone, LinkedIn stripped before LLM call")),
    ],
    col_widths=[0.55, 2.4, 0.8, 2.75],
)

h2("11.2  Active Findings — Must Fix Before Production", color=RED)
tbl(
    ["#", "Finding", "Risk", "Fix Required"],
    [
        ("C-01",
         "AI provider set to Groq/Gemini/NVIDIA in dev",
         "Candidate PII sent to US commercial servers",
         "Switch to Azure OpenAI in Settings page — see Section 7"),
        ("C-03",
         "No HTTPS — traffic over plain HTTP port 8000",
         "Passwords and resume data readable on LAN",
         "Deploy nginx + self-signed cert + GPO push — see Section 9"),
        ("C-05",
         "CORS allow_origins=['*']",
         "Any network origin can call the API",
         "Restrict to https://192.168.1.50 after SSL setup — Section 9.3"),
        ("C-07",
         "APP_ENV not set — /docs and /redoc accessible",
         "Full API documentation exposed to anyone on the LAN",
         "Set APP_ENV=production in .env on the server laptop"),
        ("C-08",
         "Candidate data isolation missing — no uploaded_by field",
         "All 100 recruiters can see all other recruiters' candidates (PII scope too wide)",
         "Add uploaded_by FK to candidates table; filter results by user role in queries"),
        ("C-09",
         "Windows port 8000 open — FastAPI directly reachable",
         "Recruiters can bypass nginx/SSL and access API over plain HTTP",
         "Windows Firewall: block inbound port 8000 from all except localhost"),
        ("H-03",
         "PDF files in /uploads have no OS access restriction",
         "Any OS user on the server laptop can read candidate resumes",
         "Set NTFS permissions on /uploads to service account only"),
        ("H-04",
         "Groq/Gemini/NVIDIA API keys in ai_config.json for production",
         "Risk of accidental re-activation of non-approved provider",
         "Delete all non-Azure keys before production go-live"),
    ],
    col_widths=[0.45, 1.8, 1.65, 2.6],
)

h2("11.3  New Finding — Backup Strategy Missing", color=RED)
callout([
    "FINDING: No backup strategy documented or implemented.",
    "",
    "RISK: The server laptop is a single point of failure.",
    "  • Laptop hardware failure → all candidate data, resume PDFs, and analysis history LOST",
    "  • Accidental file deletion → no recovery",
    "  • Theft or damage → permanent data loss",
    "",
    "FIX — Three-layer backup (see Section 20 for full procedure):",
    "  1. Automated daily database backup to a network share or OneDrive",
    "  2. Weekly copy of /uploads folder (resume PDFs)",
    "  3. Test restore quarterly — a backup never tested is not a backup",
], bg="FEF2F2", border="EF4444")

h2("11.4  New Finding — Audit Log Rotation Missing", color=AMBER)
callout([
    "FINDING: audit.log grows indefinitely — no size limit, no rotation, no retention policy.",
    "",
    "RISK: Over time (100 recruiters × dozens of requests/day) audit.log can grow to",
    "gigabytes and eventually fill the server laptop's disk, crashing the application.",
    "",
    "FIX:",
    "  Replace FileHandler in main.py with RotatingFileHandler:",
    "    maxBytes=10*1024*1024 (10 MB per file), backupCount=10 (keep last 10 files)",
    "  This caps audit logs at ~100 MB total and auto-rotates.",
    "  Retention policy: keep 30 days of logs minimum for compliance audit trail.",
], bg="FEF9C3", border="F59E0B")

h2("11.5  Backlog Findings — Not Blocking, Address Within 90 Days", color=INDIGO)
tbl(
    ["#", "Finding", "Risk", "Planned Fix"],
    [
        ("BL-01",
         "No password minimum length enforcement",
         "Weak passwords (e.g. '123') accepted by the system",
         "Add len(password) >= 8 check in users router create_user and update_user"),
        ("BL-02",
         "No account lifecycle policy — no offboarding process documented",
         "Ex-employees may retain active accounts after leaving Kforce",
         "Document: admin sets is_active=False when recruiter leaves. Add to HR offboarding checklist."),
        ("BL-03",
         "BitLocker not confirmed on server laptop",
         "If laptop is lost, database file (cms.db) and /uploads PDFs readable without password",
         "IT enables BitLocker on server laptop. Also enables PostgreSQL pg_tde when migrating DB."),
    ],
    col_widths=[0.55, 1.9, 1.8, 2.25],
)

# =============================================================================
#  SECTION 12 — REGULATORY COMPLIANCE
# =============================================================================
h1("12.  Regulatory Compliance Assessment")

tbl(
    ["Regulation", "Applicability", "Status — Production (Azure)", "Action Required"],
    [
        ("India DPDP Act 2023",
         "Applies — Indian candidates' personal data",
         OK("Compliant — Central India region keeps data in India"),
         "Add recruiter data processing acknowledgment at upload"),
        ("GDPR (EU candidates)",
         "If any EU-resident candidates are processed",
         OK("Compliant — Azure West Europe region available"),
         "Select West Europe Azure region if processing EU candidates"),
        ("Microsoft DPA",
         "Governs Azure OpenAI data handling",
         OK("Covered — same terms as M365 Copilot"),
         "No action — automatically applies to Kforce Azure subscription"),
        ("Kforce Internal Data Policy",
         "Internal HR tool — internal policy applies",
         WRN("Pending formal InfoSec review"),
         "Submit this report to Kforce InfoSec for sign-off"),
        ("ISO 27001 / SOC 2",
         "Azure OpenAI infrastructure is certified",
         WRN("Azure certified; K Recruit app needs internal review"),
         "Conduct internal security review using Section 11 findings"),
    ],
    col_widths=[1.5, 1.5, 1.7, 1.8],
)

# =============================================================================
#  SECTION 13 — DATA MINIMISATION
# =============================================================================
h1("13.  Data Minimisation — What Is Sent to Azure OpenAI")

callout([
    "PROMPT FORMAT SENT TO AZURE OPENAI (from llm_service.py):",
    "",
    '  "You are an ATS. Here is a resume: [RESUME TEXT — email/phone/LinkedIn STRIPPED]',
    '   Here is the job description: [JD TEXT]',
    '   Analyse gaps and improvements."',
    "",
    "  PII STRIPPED (v4 fix): email, phone number, LinkedIn URL removed before LLM call.",
    "  Name, location, work history, and skills retained — needed for analysis.",
    "  Both resume text (stripped) and JD go to Kforce's own Azure OpenAI resource.",
], bg="F1F5F9", border="94A3B8")

# =============================================================================
#  SECTION 14 — PRODUCTION DEPLOYMENT CHECKLIST (UPDATED V6)
# =============================================================================
h1("14.  Production Deployment Checklist — Updated v6")

h2("14.1  Blockers — Must Complete Before Any Production Use", color=RED)
tbl(
    ["#", "Action", "Owner", "Reference", "Status"],
    [
        ("B-01", "Switch AI provider to Azure OpenAI in Settings",
         "Dev / HR Admin", "Section 7", BAD("Pending")),
        ("B-02", "Provision Azure OpenAI resource in Kforce's Azure tenant",
         "IT / Cloud Team", "Section 7.2", BAD("Pending")),
        ("B-03", "Create GPT-4o deployment and obtain endpoint + API key",
         "IT / Cloud Team", "Section 7.2", BAD("Pending")),
        ("B-04", "Set APP_ENV=production in .env on server laptop",
         "Dev Team", "Finding C-07", BAD("Pending")),
        ("B-05", "Install nginx + SSL certificate (self-signed + GPO push)",
         "IT / Dev Team", "Section 9", BAD("Pending")),
        ("B-06", "Assign static IP to server laptop (e.g. 192.168.1.50)",
         "IT Team", "Section 19", BAD("Pending")),
        ("B-07", "Restrict CORS to https://192.168.1.50 in main.py",
         "Dev Team", "Section 9.3", BAD("Pending")),
        ("B-08", "Block port 8000 in Windows Firewall on server laptop",
         "IT / Dev Team", "Finding C-09", BAD("Pending")),
        ("B-09", "Add uploaded_by FK to candidates — data isolation",
         "Dev Team", "Finding C-08", BAD("Pending")),
        ("B-10", "Set up automated daily database backup",
         "IT / Dev Team", "Section 20", BAD("Pending")),
        ("B-11", "Pre-download HuggingFace model; set HF_HUB_OFFLINE=1 in .env",
         "Dev Team", "Section 5", BAD("Pending")),
        ("B-12", "Remove all Groq/Gemini/NVIDIA keys from production config",
         "Dev Team", "Finding H-04", BAD("Pending")),
        ("B-13", "Set NTFS permissions on /uploads — service account only",
         "IT Team", "Finding H-03", BAD("Pending")),
        ("B-14", "Add audit log rotation (RotatingFileHandler, 10 MB × 10 files)",
         "Dev Team", "Finding 11.4", BAD("Pending")),
        ("B-15", "Per-request token validation on all routes",
         "Dev Team", "Finding C-06", DONE("DONE v4+v5")),
        ("B-16", "Formal InfoSec sign-off on this compliance report",
         "InfoSec Team", "This document", BAD("Pending")),
    ],
    col_widths=[0.4, 2.4, 1.0, 1.1, 0.8],
)

h2("14.2  High Priority — Within 30 Days of Go-Live", color=AMBER)
tbl(
    ["#", "Action", "Owner", "Status"],
    [
        ("H-01", "Audit logging middleware",           "Dev Team", DONE("DONE v4")),
        ("H-02", "JWT 24h token expiry",               "Dev Team", DONE("DONE v5")),
        ("H-03", "Restrict /uploads OS permissions",   "IT Team",  WRN("Pending")),
        ("H-04", "Self-host React/Babel JS files",     "Dev Team", WRN("Mitigated — SRI hashes added")),
        ("H-05", "Install PostgreSQL on server laptop","IT Team",  WRN("Pending — required for 100 users")),
        ("H-06", "Disable sleep/hibernate on laptop",  "IT Team",  WRN("Pending")),
    ],
    col_widths=[0.4, 2.8, 1.1, 2.0],
)

h2("14.3  Backlog — Address Within 90 Days", color=INDIGO)
tbl(
    ["#", "Action", "Notes"],
    [
        ("BL-01", "Enforce minimum 8-character password",
         "One-line check in users router — create_user and update_user"),
        ("BL-02", "Document account offboarding procedure",
         "Admin sets is_active=False in User Management when recruiter leaves"),
        ("BL-03", "Enable BitLocker on server laptop",
         "IT enables full-disk encryption. Also pg_tde when migrating to PostgreSQL."),
        ("BL-04", "Add data retention policy — auto-purge records older than 90 days",
         "Configurable setting in the app"),
        ("BL-05", "Add recruiter consent acknowledgment at resume upload",
         "Checkbox: 'I confirm this data is processed under Kforce data policy'"),
        ("BL-06", "Set Azure Cost Management budget alert at $80/month",
         "5 minutes to configure in Azure portal"),
    ],
    col_widths=[0.4, 2.5, 3.6],
)

# =============================================================================
#  SECTION 15 — FILES REVIEWED
# =============================================================================
h1("15.  Scope of Code Review — Files Audited (v6)")

tbl(
    ["File", "What Was Reviewed"],
    [
        ("backend/main.py",                       "load_dotenv order, APP_ENV gate, CORS, audit middleware"),
        ("backend/deps.py",                        "JWT auth — SECRET_KEY load, get_current_user, require_admin"),
        ("backend/routers/auth.py",                "Login, bcrypt verify, JWT token, rate limit, /verify, /me"),
        ("backend/routers/users.py",               "User CRUD — password validation gap identified (BL-01)"),
        ("backend/routers/roles.py",               "Role CRUD — super_admin gated"),
        ("backend/routers/bot.py",                 "Auth added v5 — confirmed"),
        ("backend/routers/candidates.py",          "All endpoints — data isolation gap identified (C-08)"),
        ("backend/routers/settings.py",            "Admin gated — confirmed"),
        ("backend/routers/client_requirements.py", "Admin gated for write — confirmed"),
        ("backend/services/llm_service.py",        "PII stripping confirmed — email/phone/LinkedIn removed"),
        ("backend/services/semantic_service.py",   "One-time HuggingFace download only — confirmed"),
        ("backend/database.py",                    "Absolute SQLite path, seed_roles, seed_admin"),
        ("backend/models/candidate.py",            "No uploaded_by FK — data isolation gap C-08 confirmed"),
        ("backend/models/user.py",                 "bcrypt hash confirmed"),
        ("backend/models/role.py",                 "Permission JSON confirmed"),
        ("backend/schemas/client_requirement.py",  "from_attributes — Pydantic v2 compatible"),
        ("backend/.env",                           "APP_ENV missing (C-07), HF_HUB_OFFLINE missing (B-11)"),
        ("backend/requirements.txt",               "All dependencies audited — no unexpected network libs"),
        ("backend/static/index.html",              "SRI hashes confirmed on all 3 CDN scripts"),
        ("backend/static/preview.html",            "SRI hashes confirmed on all 3 CDN scripts"),
    ],
    col_widths=[2.5, 4.0],
)

# =============================================================================
#  SECTION 16 — QUICK REFERENCE
# =============================================================================
h1("16.  Quick Reference — Key Answers for the Compliance Team")

h2("Q1  What is the production server?")
callout([
    "A dedicated office laptop — permanently powered, sleep disabled, Kforce India office LAN.",
    "Static LAN IP assigned by IT (e.g. 192.168.1.50).",
    "Minimum spec: 16 GB RAM, 512 GB SSD.",
    "AI processing (GPT-4o) runs on Azure — the laptop handles web serving and database only.",
    "Full setup guide: Section 19.",
], bg="F0FDF4", border="16A34A")

h2("Q2  What is the authentication model?")
callout([
    "POST /auth/login → bcrypt verify → JWT token (24h expiry)",
    "All API routes require Bearer token — 401 on missing/expired/invalid",
    "Roles: super_admin (full) | admin (no role mgmt) | recruiter (own candidates)",
    "Brute force: 5 failed attempts per IP per 60s → HTTP 429",
    "Audit: every request logged — IP, user_id, method, path, status, duration",
], bg="EEF2FF", border="6366F1")

h2("Q3  What Azure services do we need and what do they cost?")
callout([
    "No new subscription — use existing Kforce Azure tenant (same as M365/Copilot)",
    "Enable Azure OpenAI → Create resource (Central India) → Deploy GPT-4o → Copy key",
    "Cost: ~$3–$10/month for 20 recruiters  |  ~$30–$80/month for 100 recruiters",
    "Pay-per-use — no standing fee. Full setup guide: Section 7.2",
], bg="EEF2FF", border="6366F1")

h2("Q4  Does K Recruit use a vector database?")
callout([
    "NO — K Recruit does not use any vector database.",
    "Names like 'pinecone', 'chroma', 'faiss' in code are skill keyword lists only.",
    "All semantic scoring uses in-memory cosine similarity — discarded after each call.",
    "Data storage: SQLite or PostgreSQL (relational) ONLY.",
], bg="F0FDF4", border="16A34A")

h2("Q5  What is the biggest outstanding compliance risk?")
callout([
    "TWO EQUAL CONCERNS:",
    "",
    "1. AI provider not yet switched to Azure (C-01)",
    "   Candidate PII currently going to Groq/Gemini/NVIDIA in development.",
    "   Fix: IT provisions Azure OpenAI → Settings page change → resolved.",
    "",
    "2. Candidate data isolation missing (C-08)",
    "   All 100 recruiters can see all candidates — no per-recruiter filtering.",
    "   Fix: Dev adds uploaded_by FK to candidates table + filter by user role.",
    "   This is a code change required before 100-user rollout.",
], bg="FEF2F2", border="EF4444")

# =============================================================================
#  SECTION 17 — V4 SECURITY HARDENING DETAILS
# =============================================================================
h1("17.  v4 Security Hardening — 9 Code-Level Fixes Applied")

tbl(
    ["Fix #", "Finding Fixed", "File Changed", "What Was Done"],
    [
        ("V4-F1", "Missing auth on all API routes",         "candidates.py, client_requirements.py, settings.py, recruiters.py", "get_current_user Dependency added to every protected endpoint"),
        ("V4-F2", "Hardcoded JWT secret",                   "deps.py",            "JWT_SECRET from env var; server refuses to start without it"),
        ("V4-F3", "/docs and /redoc in production",         "main.py",            "Hidden when APP_ENV=production (APP_ENV gap documented in v6 C-07)"),
        ("V4-F4", "No brute-force protection",              "routers/auth.py",    "5 failed attempts per IP per 60s → HTTP 429"),
        ("V4-F5", "No audit logging",                       "main.py",            "Middleware logs every request to audit.log"),
        ("V4-F6", "CDN scripts without integrity check",    "index.html, preview.html", "SRI sha384 hashes on all 3 CDN script tags"),
        ("V4-F7", "PII in AI prompts",                      "llm_service.py",     "Email, phone, LinkedIn stripped before LLM call"),
        ("V4-F8", "Plain-text login password",              "database.py, auth.py","Passwords bcrypt-hashed in users table"),
        ("V4-F9", "No session token expiry",                "routers/auth.py",    "JWT tokens include exp claim — 24 hours"),
    ],
    col_widths=[0.55, 1.8, 1.8, 2.35],
)

# =============================================================================
#  SECTION 18 — V5 ADDITIONAL FIXES
# =============================================================================
h1("18.  v5 Additional Fixes — 5 Code-Level Fixes Applied")

tbl(
    ["Fix #", "Issue", "File", "Fix Applied"],
    [
        ("V5-F1", "load_dotenv() after imports — JWT_SECRET not loaded before deps.py check",
         "main.py", "load_dotenv() moved to lines 1–2 before all imports"),
        ("V5-F2", "DATABASE_URL relative path in .env — SQLite could not open file",
         ".env", "DATABASE_URL removed; SQLite uses absolute path logic in database.py"),
        ("V5-F3", "Pydantic orm_mode UserWarning on every startup",
         "schemas/client_requirement.py", "orm_mode renamed to from_attributes"),
        ("V5-F4", "/bot/chat endpoint had no authentication",
         "routers/bot.py", "get_current_user dependency added to bot_chat"),
        ("V5-F5", "Multi-user RBAC not connected (startup blocked by V5-F1)",
         "main.py, database.py, all routers", "With V5-F1 resolved: full RBAC starts cleanly — users, roles, JWT end-to-end"),
    ],
    col_widths=[0.55, 2.0, 1.5, 2.45],
)

# =============================================================================
#  SECTION 19 — SERVER LAPTOP DEPLOYMENT GUIDE
# =============================================================================
h1("19.  Server Laptop Deployment Guide")

body(
    "This section documents the one-time setup required to configure the dedicated "
    "office laptop as the K Recruit production server for 100 recruiters."
)

h2("19.1  Minimum Hardware Requirements")
tbl(
    ["Component", "Minimum", "Reason"],
    [
        ("RAM",      "16 GB",     "Web server + PostgreSQL + semantic model loaded in memory"),
        ("Storage",  "512 GB SSD","PostgreSQL DB + /uploads resume PDFs grow over time"),
        ("OS",       "Windows 10/11 Pro or Ubuntu 20.04+", "Required for PostgreSQL + nginx"),
        ("Network",  "Ethernet (preferred) or WiFi", "Stable connection — do not let laptop sleep"),
        ("Power",    "Always plugged in", "Server must not run on battery — disable power saving"),
    ],
    col_widths=[1.2, 1.8, 3.5],
)

h2("19.2  One-Time Setup Checklist for IT")
tbl(
    ["Step", "Action", "Command / Location", "Who"],
    [
        ("1",  "Disable sleep and hibernate permanently",
         "Control Panel → Power Options → Never → Apply", "IT"),
        ("2",  "Assign static IP to the laptop",
         "Router DHCP reservation OR Windows: Network Adapter → IPv4 → Manual IP", "IT"),
        ("3",  "Install PostgreSQL 15+",
         "postgresql.org/download → install → create DB 'krecruitdb' and user 'krecruit'", "IT / Dev"),
        ("4",  "Set DATABASE_URL in .env",
         "DATABASE_URL=postgresql://krecruit:password@localhost/krecruitdb", "Dev"),
        ("5",  "Run database migration",
         "python main.py (app auto-creates all tables on first start)", "Dev"),
        ("6",  "Pre-download HuggingFace model",
         "python -c \"from sentence_transformers import SentenceTransformer; SentenceTransformer('all-MiniLM-L6-v2')\"", "Dev"),
        ("7",  "Set HF_HUB_OFFLINE=1 in .env",
         "backend/.env → add HF_HUB_OFFLINE=1", "Dev"),
        ("8",  "Set APP_ENV=production in .env",
         "backend/.env → APP_ENV=production", "Dev"),
        ("9",  "Install nginx",
         "nginx.org/download (Windows) or  apt install nginx (Linux)", "IT / Dev"),
        ("10", "Generate SSL certificate and configure nginx",
         "See Section 9 — full commands provided", "IT / Dev"),
        ("11", "Block port 8000 in Windows Firewall",
         "Windows Defender Firewall → Inbound Rules → New Rule → Port 8000 → Block", "IT"),
        ("12", "IT pushes SSL cert to all recruiter PCs via GPO",
         "See Section 9.1", "IT"),
        ("13", "Configure auto-start on boot",
         "Windows: Task Scheduler → new task → trigger At startup → run uvicorn + nginx", "IT / Dev"),
        ("14", "Set NTFS permissions on /uploads folder",
         "Right-click /uploads → Properties → Security → restrict to service account only", "IT"),
        ("15", "Test with 2–3 recruiters before full rollout",
         "Verify login, upload, analysis, download on real recruiter PCs", "Dev / HR"),
    ],
    col_widths=[0.35, 2.0, 2.7, 0.8],
)

h2("19.3  Environment Variables on Server Laptop (.env)")
callout([
    "Complete .env for the server laptop (production):",
    "",
    "  # Database — PostgreSQL for 100 users",
    "  DATABASE_URL=postgresql://krecruit:password@localhost/krecruitdb",
    "",
    "  # App",
    "  APP_HOST=0.0.0.0",
    "  APP_PORT=8000",
    "  APP_ENV=production          ← disables /docs and /redoc",
    "",
    "  # Semantic model — offline after pre-download",
    "  HF_HUB_OFFLINE=1            ← blocks all HuggingFace calls",
    "",
    "  # Auth — generate a new secret for production",
    "  JWT_SECRET=<run: python -c 'import secrets; print(secrets.token_hex(32))'>",
    "",
    "  # Azure OpenAI (fill in after IT provisions)",
    "  # Configured via Settings page in the app — not needed in .env",
], bg="F1F5F9", border="94A3B8")

# =============================================================================
#  SECTION 20 — OPERATIONAL PROCEDURES
# =============================================================================
h1("20.  Operational Procedures — Backup, Logs, Firewall")

h2("20.1  Backup Strategy")
callout([
    "RISK: Server laptop is a single point of failure.",
    "If the laptop fails or is damaged, all candidate data and resume PDFs are lost.",
    "",
    "REQUIRED: Three-layer backup",
    "",
    "Layer 1 — Daily database backup (automated):",
    "  Windows Task Scheduler runs every night at 11 PM:",
    "  pg_dump krecruitdb > C:\\Backups\\krecruitdb_YYYYMMDD.sql",
    "  Copy to network share: \\\\kforce-fileserver\\k-recruit-backups\\",
    "  Keep last 30 days. Total size: ~10–50 MB per day.",
    "",
    "Layer 2 — Weekly /uploads folder backup:",
    "  Copy backend/uploads/ to network share weekly.",
    "  Resume PDFs are not in the database — must be backed up separately.",
    "",
    "Layer 3 — Quarterly restore test:",
    "  Spin up a test laptop, restore from backup, verify app starts and data is intact.",
    "  A backup never tested is not a backup.",
], bg="FEF2F2", border="EF4444")

h2("20.2  Audit Log Rotation")
callout([
    "CURRENT STATE: audit.log grows indefinitely. Fix required before production.",
    "",
    "FIX — Replace FileHandler in main.py with RotatingFileHandler:",
    "",
    "  from logging.handlers import RotatingFileHandler",
    "  _fh = RotatingFileHandler(",
    "      _audit_log_path, maxBytes=10*1024*1024, backupCount=10, encoding='utf-8'",
    "  )",
    "",
    "  This keeps: audit.log (current) + audit.log.1 through audit.log.10",
    "  Total max size: ~100 MB",
    "  Oldest logs auto-deleted when limit reached",
    "",
    "  Retention policy: keep minimum 30 days of logs for compliance audit trail.",
], bg="FEF9C3", border="F59E0B")

h2("20.3  Windows Firewall Rules")
tbl(
    ["Rule", "Direction", "Port", "Action", "Reason"],
    [
        ("Block direct FastAPI access", "Inbound", "8000", BAD("BLOCK from all except localhost"), "Force all traffic through nginx/SSL"),
        ("Allow HTTPS",                 "Inbound", "443",  OK("ALLOW from LAN"),                   "Recruiter access via nginx"),
        ("Allow HTTP redirect",         "Inbound", "80",   OK("ALLOW from LAN"),                   "Redirect to HTTPS"),
        ("Allow PostgreSQL",            "Inbound", "5432", BAD("BLOCK from all except localhost"),  "DB must not be reachable from LAN"),
        ("Allow Azure OpenAI",          "Outbound","443",  OK("ALLOW to *.openai.azure.com"),       "AI provider calls"),
    ],
    col_widths=[1.9, 0.85, 0.6, 1.6, 2.55],
)

h2("20.4  Account Lifecycle (Backlog BL-02)")
callout([
    "WHEN A RECRUITER JOINS:",
    "  Admin logs into K Recruit → User Management → Create User",
    "  Set name, email, temporary password, role = recruiter",
    "  Recruiter logs in and changes password on first use (manual step until enforced)",
    "",
    "WHEN A RECRUITER LEAVES KFORCE:",
    "  HR notifies admin on the last working day",
    "  Admin → User Management → Edit User → set is_active = False",
    "  Account immediately blocked — JWT tokens rejected on next request",
    "  Do NOT delete the account — retain for audit trail",
    "",
    "NOTE: This procedure should be added to Kforce HR offboarding checklist.",
], bg="EEF2FF", border="6366F1")

# =============================================================================
#  SECTION 21 — DECLARATION
# =============================================================================
h1("21.  Declaration")

callout([
    "FOR THE KFORCE US COMPLIANCE AND IT GOVERNANCE TEAM:",
    "",
    "  K Recruit is approved for production deployment on the dedicated office laptop",
    "  subject to completing the blockers listed in Section 14.1 (B-01 through B-14, B-16).",
    "  Blocker B-15 (per-request token validation) is COMPLETE.",
    "",
    "  SECURITY STATUS AFTER v5 + v6 REVIEW:",
    "  ✓  All API endpoints require valid JWT authentication",
    "  ✓  Passwords stored as bcrypt hashes — no plain text credentials",
    "  ✓  JWT tokens expire after 24 hours",
    "  ✓  Brute-force login protection active",
    "  ✓  Full audit log of every HTTP request",
    "  ✓  SRI hashes on all CDN script tags",
    "  ✓  PII stripped from AI prompts",
    "  ✓  Multi-user RBAC: super_admin / admin / recruiter",
    "",
    "  OUTSTANDING BLOCKERS (require action before go-live):",
    "  ✗  Switch AI provider to Azure OpenAI",
    "  ✗  Deploy nginx + SSL on server laptop",
    "  ✗  Set APP_ENV=production and block port 8000",
    "  ✗  Add candidate data isolation (uploaded_by FK)",
    "  ✗  Set up daily database backup",
    "  ✗  Add audit log rotation",
    "",
    "  This document (v6) is the complete and final compliance reference.",
    "  It fully supersedes v1–v5 (2026-001 through 2026-005).",
], bg="EEF2FF", border="6366F1")

sp(8)
for label, value in [
    ("Application Name: ", "K Recruit — HR Resume Optimizer"),
    ("Report Version: ",   "6.0 — Final  (Server Laptop & Operational Gaps Edition)"),
    ("Document Ref: ",     "K-RECRUIT-COMP-2026-006"),
    ("Supersedes: ",       "v1–v5: CMS-TalentAI-COMP-2026-001 through K-RECRUIT-COMP-2026-005"),
    ("Prepared by: ",      "Technical Architecture Review, K Recruit Development Team"),
    ("Contact: ",          "sushrut.nistane@kforce.co.in"),
    ("Date: ",             "May 2026"),
]:
    p = doc.add_paragraph()
    rn(p, label, bold=True,  sz=9, color=NAVY)
    rn(p, value, bold=False, sz=9, color=MUTED)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)

# ── Save ──────────────────────────────────────────────────────────────────────
out = r"C:\Users\SushrutNistane\OneDrive - Kforce Inter\cms-resume-optimizer\K_Recruit_Compliance_Report_v6.docx"
doc.save(out)
print(f"Saved: {out}")
