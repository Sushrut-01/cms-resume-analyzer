"""Generate CMS TalentAI Compliance Report VERSION 2 as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY    = RGBColor(0x0F, 0x17, 0x2A)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
MUTED   = RGBColor(0x64, 0x74, 0x8B)
TEAL    = RGBColor(0x0D, 0x6E, 0x6E)
RED     = RGBColor(0xC0, 0x10, 0x20)
GREEN   = RGBColor(0x16, 0x65, 0x34)
AMBER   = RGBColor(0x92, 0x40, 0x0E)
INDIGO  = RGBColor(0x44, 0x38, 0xCA)
GRAY_L  = RGBColor(0x94, 0xA3, 0xB8)
INDIGO_L= RGBColor(0x99, 0xA3, 0xFF)

# ── XML helpers ───────────────────────────────────────────────────────────────
def shade_cell(cell, hex6):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex6.lstrip("#"))
    tcPr.append(shd)

def cell_padding(cell, twips=160):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ("top","left","bottom","right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(twips))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

def add_bottom_border(para, color="0F172A", sz="6"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    sz)
    bot.set(qn("w:space"), "2")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def left_border_cell(cell, color="6366F1", sz="18"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBorders")
    for edge in ("top","bottom","right"):
        t = OxmlElement(f"w:{edge}")
        t.set(qn("w:val"), "none"); tcBdr.append(t)
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    sz)
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), color)
    tcBdr.append(left)
    tcPr.append(tcBdr)

# ── Typography helpers ────────────────────────────────────────────────────────
def run(para, text, bold=False, italic=False, size=10, color=None, underline=False):
    r = para.add_run(text)
    r.bold      = bold
    r.italic    = italic
    r.underline = underline
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return r

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
    add_bottom_border(p)
    return p

def h2(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11.5)
    r.font.color.rgb = color or NAVY
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = MUTED
    return p

def body(text, size=10, color=None, space_after=5, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = color or NAVY
    return p

def bullet(text, size=10, indent=0.35):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(indent)
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = NAVY
    return p

def spacer(pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)

def callout(lines, bg="EEF2FF", border="6366F1"):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0,0)
    shade_cell(cell, bg)
    left_border_cell(cell, border, sz="18")
    cell_padding(cell, 200)
    cell.paragraphs[0].clear()
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i==0 else cell.add_paragraph()
        r = p.add_run(line)
        r.font.size = Pt(9.5); r.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
    spacer(6)

def table(headers, rows, col_widths=None, hdr_bg="0F172A", alt="F8FAFC"):
    nc  = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=nc)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    hr = tbl.rows[0]
    for j,h in enumerate(headers):
        c = hr.cells[j]; shade_cell(c, hdr_bg); cell_padding(c, 120)
        p = c.paragraphs[0]; r = p.add_run(h)
        r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
    # Rows
    for i,row in enumerate(rows):
        bg = alt if i%2==0 else "FFFFFF"
        tr = tbl.rows[i+1]
        for j, cv in enumerate(row):
            c = tr.cells[j]; shade_cell(c, bg); cell_padding(c, 120)
            p = c.paragraphs[0]
            if isinstance(cv, tuple):
                text, bold, clr = cv
                r = p.add_run(str(text)); r.bold = bold
                r.font.size = Pt(9); r.font.color.rgb = clr
            else:
                r = p.add_run(str(cv))
                r.font.size = Pt(9); r.font.color.rgb = NAVY
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
    if col_widths:
        for j,w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[j].width = Inches(w)
    spacer(6)

# ── Shorthand tuples ──────────────────────────────────────────────────────────
def OK(t):  return (t, True,  GREEN)
def BAD(t): return (t, True,  RED)
def WRN(t): return (t, True,  AMBER)
def INF(t): return (t, False, MUTED)
def BLU(t): return (t, True,  INDIGO)

# =============================================================================
#  COVER
# =============================================================================
cover_tbl  = doc.add_table(rows=1, cols=1)
cover_cell = cover_tbl.cell(0,0)
shade_cell(cover_cell, "0F172A")
cell_padding(cover_cell, 360)
cover_cell.paragraphs[0].clear()

def cl(text, sz, clr, bold=False, after=6):
    p = cover_cell.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(sz); r.font.color.rgb = clr
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(after)
    return p

cl("CMS TalentAI",                                      24, WHITE,    bold=True,  after=2)
cl("HR Resume Optimizer",                               14, GRAY_L,   bold=False, after=18)
cl("DATA PRIVACY & COMPLIANCE ASSESSMENT REPORT",      12, INDIGO_L, bold=True,  after=4)
cl("VERSION 2  —  PRODUCTION DEPLOYMENT EDITION",       9, GRAY_L,   bold=False, after=18)

for lbl, val in [
    ("Document Ref:",    "CMS-TalentAI-COMP-2026-002"),
    ("Version:",         "2.0  |  Supersedes v1.0 (CMS-TalentAI-COMP-2026-001)"),
    ("Prepared for:",    "Kforce US Compliance / IT Governance Team"),
    ("Prepared by:",     "Application Compliance Review — Technical Architecture"),
    ("Scope:",           "Full library audit + Production-only Azure OpenAI architecture"),
    ("Review Date:",     "May 2026"),
    ("Classification:",  "Internal Confidential"),
]:
    p = cover_cell.add_paragraph()
    r1 = p.add_run(f"{lbl:20s}"); r1.bold = True
    r1.font.size = Pt(9); r1.font.color.rgb = GRAY_L
    r2 = p.add_run(val)
    r2.font.size = Pt(9); r2.font.color.rgb = WHITE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)

cover_cell.paragraphs[0]._element.getparent().remove(
    cover_cell.paragraphs[0]._element)
spacer(12)

# =============================================================================
#  SECTION 1 — SCOPE AND CHANGES FROM V1
# =============================================================================
h1("1.  What Is New in Version 2")

callout([
    "VERSION 2 SCOPE CHANGES vs Version 1:",
    "",
    "  1. Production AI provider confirmed as Azure OpenAI (Microsoft Copilot) ONLY.",
    "     All other providers (Groq, Gemini, NVIDIA NIM, Ollama) are EXCLUDED from production.",
    "",
    "  2. Complete library-by-library network audit added — every Python package and",
    "     every frontend dependency reviewed for external data transmission.",
    "",
    "  3. Sentence-transformers model download behaviour documented and remediated.",
    "",
    "  4. Frontend CDN dependency risk assessed and offline mitigation provided.",
    "",
    "  5. Production-ready deployment checklist updated to reflect Azure-only posture.",
], bg="EEF2FF", border="6366F1")

# =============================================================================
#  SECTION 2 — PRODUCTION AI PROVIDER DECISION
# =============================================================================
h1("2.  Production AI Provider Decision")

h2("2.1  Confirmed Production Architecture")
body(
    "Based on the organisation's requirement that no candidate data leaves Kforce's controlled "
    "data boundary, the production deployment will use ONLY Microsoft Azure OpenAI Service — "
    "the same AI infrastructure that powers Microsoft 365 Copilot. All other AI providers "
    "are permanently excluded from the production configuration."
)

table(
    ["Provider", "Development Use?", "Production Use?", "Reason"],
    [
        ("Azure OpenAI (Copilot)",  OK("✓ Allowed"),  OK("✓ APPROVED"),   "Within Kforce's Microsoft tenant — Microsoft DPA covers all data"),
        ("Ollama (Local LLM)",      OK("✓ Allowed"),  WRN("Optional only"),"100% local — zero egress. Acceptable if Azure is unavailable temporarily"),
        ("Groq Cloud",              WRN("Dev only"),  BAD("✗ BLOCKED"),    "Third-party US server. PII sent outside organisation boundary"),
        ("Google Gemini",           WRN("Dev only"),  BAD("✗ BLOCKED"),    "Third-party Google server. Data retention policy not enterprise-grade"),
        ("NVIDIA NIM",              WRN("Dev only"),  BAD("✗ BLOCKED"),    "Third-party NVIDIA server. Research-use data policy insufficient"),
    ],
    col_widths=[1.6, 1.1, 1.2, 2.6],
)

h2("2.2  Why Azure OpenAI Is Compliant for Production")
for item in [
    "Azure OpenAI is provisioned inside Kforce's own Azure subscription — Kforce is the data controller",
    "Microsoft's Data Processing Addendum (DPA) applies — same as all M365 / Copilot services",
    "Microsoft commits that prompts and completions are NOT used to train OpenAI models (enterprise setting)",
    "Data residency is configurable — choose India or US region per your Azure subscription geography",
    "Azure OpenAI has ISO 27001, SOC 2 Type II, and GDPR Article 28 certifications",
    "The connector is already built in the application — it is a settings change, not a code change",
]:
    bullet(item)

h2("2.3  How to Enable Azure OpenAI in the Application")
table(
    ["Step", "Action", "Who"],
    [
        ("1", "Raise Azure OpenAI resource request in Kforce's Azure portal (portal.azure.com → Azure OpenAI)", "IT / Cloud Team"),
        ("2", "Create a GPT-4o deployment within that resource", "IT / Cloud Team"),
        ("3", "Copy the endpoint URL (https://YOUR-RESOURCE.openai.azure.com) and API key", "IT / Cloud Team"),
        ("4", "Open CMS TalentAI → Settings → select Azure OpenAI → enter endpoint + key → Save", "Dev / HR Team"),
        ("5", "Click 'Test Connection' in Settings to confirm it is working", "Dev Team"),
        ("6", "Remove Groq / Gemini keys from ai_config.json (or move to env variables)", "Dev Team"),
    ],
    col_widths=[0.35, 4.3, 1.5],
)

# =============================================================================
#  SECTION 3 — COMPLETE LIBRARY NETWORK AUDIT
# =============================================================================
h1("3.  Complete Library-by-Library Network Audit")

body(
    "This section documents every Python library and every frontend dependency used by the application. "
    "Each was reviewed for: (a) does it make outbound network calls? (b) does it transmit user data? "
    "(c) when does any network call occur? This is a fact-based audit of the actual source code."
)

# ── 3.1 Backend Python libraries ─────────────────────────────────────────────
h2("3.1  Backend Python Libraries (requirements.txt)")

table(
    ["Library", "Version", "Purpose", "Makes Network Calls?", "Transmits User Data?", "Verdict"],
    [
        ("fastapi",          "≥0.110",  "Web API framework",
         OK("No — runs locally"),        OK("No"),  OK("SAFE")),
        ("uvicorn",          "≥0.29",   "ASGI web server",
         OK("No — runs locally"),        OK("No"),  OK("SAFE")),
        ("python-multipart", "≥0.0.9",  "File upload parsing",
         OK("No — runs locally"),        OK("No"),  OK("SAFE")),
        ("sqlalchemy",       "≥2.0",    "Database ORM (SQLite/PostgreSQL)",
         OK("No — local DB only"),       OK("No"),  OK("SAFE")),
        ("psycopg2-binary",  "≥2.9",    "PostgreSQL database driver",
         OK("No — local DB only"),       OK("No"),  OK("SAFE")),
        ("pymupdf (fitz)",   "≥1.24",   "PDF text extraction",
         OK("No — runs locally"),        OK("No"),  OK("SAFE")),
        ("python-docx",      "≥1.1",    "Word document generation",
         OK("No — runs locally"),        OK("No"),  OK("SAFE")),
        ("requests",         "≥2.31",   "HTTP client — calls AI provider API only",
         WRN("YES — calls Azure OpenAI"),WRN("Resume + JD text sent to Azure"),  WRN("APPROVED *")),
        ("python-dotenv",    "≥1.0",    "Reads .env file for config",
         OK("No — reads local file"),    OK("No"),  OK("SAFE")),
        ("sentence-transformers","≥2.7","Semantic similarity scoring",
         WRN("YES — one-time model download only"),OK("No — model file only"),  WRN("SEE NOTE **")),
        ("scikit-learn",     "≥1.3",    "TF-IDF fallback similarity scoring",
         OK("No — runs locally"),        OK("No"),  OK("SAFE")),
        ("numpy",            "(dep)",   "Math operations (dep of above)",
         OK("No — runs locally"),        OK("No"),  OK("SAFE")),
    ],
    col_widths=[1.55, 0.6, 1.7, 1.5, 1.3, 0.8],
)

callout([
    "*  REQUESTS LIBRARY — approved for production because:",
    "   In production configuration, 'requests' calls ONLY the Azure OpenAI endpoint",
    "   (https://YOUR-RESOURCE.openai.azure.com) which is inside Kforce's Azure tenant.",
    "   The Groq/Gemini/NVIDIA endpoints are NOT called when provider = 'azure'.",
    "   No other outbound HTTP call exists anywhere in the backend codebase.",
    "   Confirmed by tracing every requests.post() and requests.get() call in llm_service.py and settings.py.",
], bg="F0FDF4", border="16A34A")

callout([
    "**  SENTENCE-TRANSFORMERS — MODEL DOWNLOAD (important detail):",
    "",
    "    What happens: On the FIRST startup after installation, this library downloads the",
    "    'all-MiniLM-L6-v2' model (~22 MB) from HuggingFace (huggingface.co).",
    "",
    "    What is downloaded: A static neural network model file — binary weights only.",
    "    NO RESUME DATA, NO PII, and NO USER CONTENT is ever sent to HuggingFace.",
    "    This is equivalent to downloading a software installer from a vendor's website.",
    "",
    "    After first download: The model is cached locally. ALL subsequent similarity",
    "    calculations run entirely on the local server with zero network calls.",
    "",
    "    PRODUCTION MITIGATION: Pre-download the model during server setup (before go-live).",
    "    Command: python -c \"from sentence_transformers import SentenceTransformer;",
    "             SentenceTransformer('all-MiniLM-L6-v2')\"",
    "    After this, the server can run air-gapped with no HuggingFace access needed.",
], bg="FEF9C3", border="F59E0B")

# ── 3.2 Frontend dependencies ─────────────────────────────────────────────────
h2("3.2  Frontend JavaScript Dependencies (index.html / preview.html)")

body(
    "The frontend is a single-page HTML application with no build process. It loads three "
    "JavaScript libraries from a CDN on initial page load. These are reviewed below."
)

table(
    ["Library", "Source", "Version", "What is fetched?", "User Data Sent?", "Verdict"],
    [
        ("React",     "cdnjs.cloudflare.com", "18.2.0",
         "react.production.min.js (42 KB)", OK("No — static JS file"), OK("SAFE")),
        ("ReactDOM",  "cdnjs.cloudflare.com", "18.2.0",
         "react-dom.production.min.js (130 KB)", OK("No — static JS file"), OK("SAFE")),
        ("Babel",     "cdnjs.cloudflare.com", "7.23.2",
         "babel.min.js (905 KB — JSX transpiler)", OK("No — static JS file"), OK("SAFE")),
    ],
    col_widths=[0.9, 2.2, 0.7, 2.1, 1.2, 0.8],
)

callout([
    "CDN LIBRARIES — RISK ASSESSMENT:",
    "",
    "  What the CDN sees: Only a standard HTTP GET request for a JavaScript file.",
    "  Cloudflare/CDNJS receives the browser's IP address and a file request.",
    "  NO resume data, candidate PII, or application data is ever sent to the CDN.",
    "  These libraries run entirely inside the recruiter's browser after download.",
    "",
    "  PRODUCTION MITIGATION (for fully air-gapped / zero-CDN deployment):",
    "  Download the three JS files once and serve them from the local FastAPI server.",
    "  Change the three <script src> tags to point to /static/js/react.min.js etc.",
    "  This eliminates the only remaining external contact the browser makes.",
    "",
    "  RISK LEVEL WITHOUT MITIGATION: LOW — CDN sees browser IP only, no user data.",
    "  RISK LEVEL WITH MITIGATION: ZERO — fully air-gapped.",
], bg="FEF9C3", border="F59E0B")

# ── 3.3 Complete outbound call map ─────────────────────────────────────────────
h2("3.3  Complete Outbound Network Call Map — Production (Azure Only)")

body("The following is the complete and exhaustive list of ALL outbound network connections "
     "made by the production application when configured with Azure OpenAI provider:")

table(
    ["Call #", "Source File", "Function", "Destination", "Frequency", "Data Sent", "Verdict"],
    [
        ("1", "llm_service.py", "_azure_openai()",
         "YOUR-RESOURCE.openai.azure.com",
         "Per AI analysis request",
         WRN("Resume text + JD text"),
         WRN("APPROVED — within Azure tenant")),
        ("2", "settings.py",    "test_connection()",
         "YOUR-RESOURCE.openai.azure.com",
         "Manual test only (admin)",
         OK("'Hello' — 5 tokens only"),
         OK("SAFE")),
        ("3", "semantic_service.py","_load_tier1()",
         "huggingface.co",
         "ONE-TIME on first startup",
         OK("No user data — model file download"),
         WRN("MITIGATED — pre-download model")),
        ("4", "index.html (browser)", "<script src>",
         "cdnjs.cloudflare.com",
         "Once per browser session",
         OK("No user data — JS file download"),
         WRN("MITIGATED — self-host JS files")),
    ],
    col_widths=[0.3, 1.1, 1.3, 1.7, 0.9, 1.3, 1.1],
)

callout([
    "CONFIRMED: In production (Azure provider), there are ONLY 4 distinct outbound connections.",
    "Connections #1 and #2 go to Kforce's own Azure tenant — covered by Microsoft DPA.",
    "Connections #3 and #4 can be eliminated entirely with the mitigations described above.",
    "NO other outbound call exists anywhere in the codebase — confirmed by full code review.",
], bg="F0FDF4", border="16A34A")

# =============================================================================
#  SECTION 4 — DATA INVENTORY (updated)
# =============================================================================
h1("4.  Data Inventory — What Personal Data Is Processed")

table(
    ["Data Element", "Stored Where", "Sent to Azure OpenAI?", "Purpose of AI Processing"],
    [
        ("Candidate Full Name",  "Local SQLite/PG DB",
         WRN("YES — in resume text"),  "Contact completeness check in AI output"),
        ("Email Address",        "Local SQLite/PG DB",
         WRN("YES — in resume text"),  "Contact completeness check"),
        ("Mobile Number",        "Local SQLite/PG DB",
         WRN("YES — in resume text"),  "Contact completeness check"),
        ("Location / City",      "Local SQLite/PG DB",
         WRN("YES — in resume text"),  "Resume context for analysis"),
        ("LinkedIn URL",         "Local SQLite/PG DB",
         WRN("YES — in resume text"),  "Profile reference in resume"),
        ("Full Resume Text",     "Local SQLite/PG DB + /uploads",
         WRN("YES — core prompt"),     "Gap analysis, JD alignment, skill scoring"),
        ("Job Description Text", "Local SQLite/PG DB",
         WRN("YES — core prompt"),     "Comparison target for gap analysis"),
        ("AI Analysis Results",  "Local SQLite/PG DB",
         OK("No — generated locally"),  "Stored as output, never re-sent"),
        ("Login Credentials",    "ai_config.json (local)",
         OK("No"),                       "Application authentication only"),
        ("Recruiter Names",      "Local SQLite/PG DB",
         OK("No"),                       "Internal workflow tracking only"),
        ("PDF Files",            "/uploads folder (local)",
         OK("No — text extracted locally"), "PDF stored for audit trail only"),
    ],
    col_widths=[1.5, 1.5, 1.4, 2.1],
)

callout([
    "DATA MINIMISATION OPPORTUNITY (recommended for v3):",
    "Name, Email, and Mobile are extracted and stored BEFORE the AI call (by pdf_service.py).",
    "A future enhancement can strip these three fields from the AI prompt, reducing PII",
    "transmitted to Azure OpenAI while preserving 100% of the analytical functionality.",
    "This would limit the AI prompt to: work history, skills, projects, education, JD text only.",
], bg="EEF2FF", border="6366F1")

# =============================================================================
#  SECTION 5 — DATA FLOW: PRODUCTION
# =============================================================================
h1("5.  Production Data Flow Diagram")

callout([
    "PRODUCTION ARCHITECTURE — Azure OpenAI Only",
    "",
    "  ┌─────────────────────────────────────────────────────────────────┐",
    "  │                    KFORCE OFFICE PREMISES                       │",
    "  │                                                                 │",
    "  │  [HR Recruiter PC] ─── HTTPS/TLS ───► [Application Server]     │",
    "  │                                         │                       │",
    "  │  PDF Resume uploaded ──────────────────►│                       │",
    "  │                                         ├── pdf_service.py      │",
    "  │                                         │     (local parsing)   │",
    "  │                                         ├── SQLite / PostgreSQL │",
    "  │                                         │     (local storage)   │",
    "  │                                         └── /uploads (local)    │",
    "  │                                                │                │",
    "  └────────────────────────────────────────────────┼────────────────┘",
    "                                                    │ Resume text + JD",
    "                                                    │ HTTPS/443 only",
    "                                                    ▼",
    "                               ┌──────────────────────────────────┐",
    "                               │  KFORCE AZURE TENANT             │",
    "                               │  Azure OpenAI Service            │",
    "                               │  ├── GPT-4o deployment           │",
    "                               │  ├── Microsoft DPA applies       │",
    "                               │  ├── No model training           │",
    "                               │  └── Data residency configurable │",
    "                               └──────────────────────────────────┘",
    "                                         │ AI analysis result",
    "                                         │ (returned as JSON)",
    "                               ┌─────────▼──────────────┐",
    "                               │ Stored in local DB only │",
    "                               └────────────────────────┘",
], bg="F0FDF4", border="16A34A")

# =============================================================================
#  SECTION 6 — SECURITY FINDINGS
# =============================================================================
h1("6.  Security Findings & Pre-Production Checklist")

h2("6.1  Critical — Must Fix Before Go-Live", color=RED)
table(
    ["#", "Finding", "Code Location", "Risk", "Fix Required"],
    [
        ("C-01", "API keys stored in plain text in ai_config.json",
         "backend/ai_config.json",
         "All keys exposed if server is accessed",
         "Move to OS environment variables or Windows Credential Manager"),
        ("C-02", "No HTTPS — traffic over plain HTTP port 8000",
         "main.py",
         "Credentials and resume data in clear text on LAN",
         "Deploy behind nginx or IIS with TLS certificate"),
        ("C-03", "Login password stored in plain text",
         "ai_config.json: login_password",
         "Anyone with file access sees the password",
         "Hash with bcrypt; store in environment variable"),
        ("C-04", "CORS allows all origins: allow_origins=['*']",
         "main.py line 28",
         "Any origin on the network can call the API",
         "Restrict to internal server IP / intranet domain"),
        ("C-05", "No per-request authentication check on API routes",
         "All routers — candidates, settings, recruiters",
         "API endpoints accessible without valid session token",
         "Add token verification FastAPI Dependency to all protected routes"),
    ],
    col_widths=[0.35, 1.6, 1.5, 1.5, 1.55],
)

h2("6.2  High Priority", color=AMBER)
table(
    ["#", "Finding", "Risk", "Fix Required"],
    [
        ("H-01", "No audit log — no record of who accessed which candidate",
         "No traceability for compliance audit",
         "Add logging middleware: datetime, user, action, candidate_id"),
        ("H-02", "No token expiry — auth token valid indefinitely",
         "Stolen token = permanent access",
         "Add timestamp to token or switch to JWT with 8-hour expiry"),
        ("H-03", "PDF files in /uploads have no OS access restriction",
         "Any OS user can read candidate resumes",
         "chmod 700 /uploads or restrict NTFS permissions to service account"),
        ("H-04", "Groq/NVIDIA/Gemini keys still present in ai_config.json",
         "Risk of accidental activation of non-approved provider",
         "Remove all non-Azure keys from config before production deployment"),
    ],
    col_widths=[0.35, 2.0, 1.7, 2.45],
)

h2("6.3  Medium — Recommended", color=MUTED)
table(
    ["#", "Finding", "Fix Required"],
    [
        ("M-01", "Database file (cms.db) has no encryption at rest",
         "Enable SQLite encryption (SQLCipher) or use PostgreSQL with encrypted tablespace"),
        ("M-02", "No data retention policy — records stored indefinitely",
         "Add configurable auto-purge (e.g. delete records older than 90 days)"),
        ("M-03", "No candidate consent notice at upload",
         "Add recruiter acknowledgment at upload: 'I confirm this candidate's data is processed under Kforce's data policy'"),
        ("M-04", "Frontend loads React/Babel from external CDN (Cloudflare)",
         "Download JS files once and serve from /static/js/ (eliminates all CDN dependency)"),
        ("M-05", "sentence-transformers model downloaded from HuggingFace on first start",
         "Pre-download model during server setup; thereafter server runs fully air-gapped"),
    ],
    col_widths=[0.35, 2.0, 4.15],
)

# =============================================================================
#  SECTION 7 — REGULATORY COMPLIANCE
# =============================================================================
h1("7.  Regulatory Compliance Assessment")

table(
    ["Regulation", "Applicability", "Status with Azure OpenAI", "Action Required"],
    [
        ("India DPDP Act 2023",
         "Applies — Indian candidates' personal data",
         OK("Compliant — data stays in Azure India/enterprise tenant"),
         "Add recruiter data processing acknowledgment at upload"),
        ("GDPR (EU candidates)",
         "If any EU-resident candidates are processed",
         OK("Compliant — Azure OpenAI with EU region available"),
         "Select EU Azure region if processing EU candidates"),
        ("Microsoft DPA",
         "Governs Azure OpenAI data handling",
         OK("Covered — same as M365 Copilot terms"),
         "No action — automatically applies to Azure subscription"),
        ("Kforce Internal Data Policy",
         "Internal HR tool — internal policy applies",
         WRN("Pending formal InfoSec review"),
         "Submit this report to Kforce InfoSec for sign-off"),
        ("ISO 27001 / SOC 2",
         "Azure OpenAI is certified — application itself is not",
         WRN("Azure provider certified; app not yet certified"),
         "Conduct internal security review using findings in Section 6"),
    ],
    col_widths=[1.5, 1.5, 1.6, 1.9],
)

# =============================================================================
#  SECTION 8 — PRODUCTION DEPLOYMENT CHECKLIST
# =============================================================================
h1("8.  Production Deployment Checklist")

h2("8.1  Before Go-Live (Blockers)")
table(
    ["#", "Action", "Owner", "Status"],
    [
        ("1", "Provision Azure OpenAI resource in Kforce's Azure tenant", "IT / Cloud Team", BAD("Pending")),
        ("2", "Create GPT-4o deployment, obtain endpoint URL and API key", "IT / Cloud Team", BAD("Pending")),
        ("3", "Move all API keys and login password to environment variables", "Dev Team", BAD("Pending")),
        ("4", "Enable HTTPS on the server (nginx + TLS certificate)", "IT / Dev Team", BAD("Pending")),
        ("5", "Restrict CORS to internal office IP range in main.py", "Dev Team", BAD("Pending")),
        ("6", "Add per-request token validation to all API routes", "Dev Team", BAD("Pending")),
        ("7", "Remove Groq/Gemini/NVIDIA keys from production config", "Dev Team", BAD("Pending")),
        ("8", "Pre-download sentence-transformers model on server", "Dev Team", BAD("Pending")),
        ("9", "InfoSec sign-off on this compliance report", "InfoSec Team", BAD("Pending")),
    ],
    col_widths=[0.3, 3.5, 1.5, 0.9],
)

h2("8.2  Post Go-Live (Within 30 Days)")
table(
    ["#", "Action", "Owner"],
    [
        ("10", "Implement audit logging middleware (user + action + timestamp)", "Dev Team"),
        ("11", "Add session token expiry (8-hour JWT or equivalent)", "Dev Team"),
        ("12", "Restrict /uploads folder OS permissions to service account", "IT Team"),
        ("13", "Encrypt database at rest (SQLCipher or PostgreSQL TDE)", "IT / Dev Team"),
        ("14", "Self-host React/Babel JS files (eliminate CDN dependency)", "Dev Team"),
        ("15", "Add data retention policy — auto-purge records > 90 days", "Dev Team"),
    ],
    col_widths=[0.3, 4.4, 1.5],
)

h2("8.3  What Does NOT Need to Change for Compliance")
for item in [
    "Core application logic — all business features are sound and production-ready",
    "Database schema — no structural changes required",
    "Resume analysis, JD alignment, and gap detection — fully functional",
    "Authentication flow — structure is correct (needs HTTPS + token expiry only)",
    "No third-party analytics, tracking, or telemetry in the application",
    "No phone-home, auto-update, or background telemetry anywhere in the codebase",
    "PDF extraction (pymupdf) — runs 100% locally, no external calls",
    "Word document generation (python-docx) — runs 100% locally, no external calls",
    "Semantic similarity scoring (after model pre-download) — runs 100% locally",
]:
    bullet(item)

# =============================================================================
#  SECTION 9 — FILES REVIEWED
# =============================================================================
h1("9.  Scope of Code Review — Files Audited")

table(
    ["File", "What Was Reviewed"],
    [
        ("backend/main.py",                  "Entry point, CORS config, middleware, router registration"),
        ("backend/routers/candidates.py",    "All candidate endpoints: upload, analyze, align, download, approve"),
        ("backend/routers/auth.py",          "Login, token generation, token verification"),
        ("backend/routers/settings.py",      "Settings API, API key masking, connection test, domain coverage"),
        ("backend/routers/recruiters.py",    "Recruiter CRUD endpoints"),
        ("backend/routers/client_requirements.py", "Job Description CRUD endpoints"),
        ("backend/services/llm_service.py",  "ALL AI provider calls: every requests.post() traced and documented"),
        ("backend/services/pdf_service.py",  "PDF text extraction — confirmed fully local (fitz/PyMuPDF)"),
        ("backend/services/docx_service.py", "Word doc generation — confirmed fully local (python-docx)"),
        ("backend/services/semantic_service.py","Semantic scoring — network call confirmed as model download only"),
        ("backend/database.py",              "Database config — SQLite local / PostgreSQL production path"),
        ("backend/ai_config.py",             "Config management — plain text storage issue documented"),
        ("backend/ai_config.json",           "Active config — API keys, credentials, provider setting"),
        ("backend/requirements.txt",         "All Python dependencies — each audited for network behaviour"),
        ("backend/static/index.html",        "Frontend — CDN dependencies, API calls, data handling"),
        ("backend/static/preview.html",      "Premium frontend — same CDN dependencies, same API calls"),
    ],
    col_widths=[2.5, 4.0],
)

# =============================================================================
#  SECTION 10 — FINAL DECLARATION
# =============================================================================
h1("10.  Declaration")

callout([
    "FOR THE US COMPLIANCE TEAM:",
    "",
    "  This application is approved for production deployment subject to completing",
    "  the 9 blocker items listed in Section 8.1.",
    "",
    "  The single most important action is: provision Azure OpenAI within Kforce's",
    "  Azure tenant and provide the endpoint + key to the development team.",
    "  This is a settings change — no code deployment is needed.",
    "",
    "  Once Azure OpenAI is configured:",
    "  • All AI processing stays within Kforce's Microsoft data boundary",
    "  • Covered by the same Microsoft DPA as M365 Copilot",
    "  • No candidate PII reaches any third-party server",
    "  • The only external contact is to your own Azure OpenAI resource",
    "",
    "  Every library in this application has been individually audited.",
    "  No hidden telemetry, tracking, or data exfiltration was found.",
], bg="EEF2FF", border="6366F1")

spacer(6)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
for label, value in [
    ("Report Version: ",   "2.0 — Final"),
    ("Supersedes: ",       "CMS-TalentAI-COMP-2026-001 (Version 1)"),
    ("Prepared by: ",      "Technical Architecture Review, CMS TalentAI Development Team"),
    ("Contact: ",          "sushrut.nistane@kforce.co.in"),
    ("Date: ",             "May 2026"),
]:
    p = doc.add_paragraph()
    run(p, label,  bold=True,  size=9, color=NAVY)
    run(p, value,  bold=False, size=9, color=MUTED)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)

# ── Save ──────────────────────────────────────────────────────────────────────
out = r"C:\Users\SushrutNistane\OneDrive - Kforce Inter\cms-resume-optimizer\CMS_TalentAI_Compliance_Report_v2.docx"
doc.save(out)
print(f"Saved: {out}")
