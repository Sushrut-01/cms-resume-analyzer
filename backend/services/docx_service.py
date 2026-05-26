# =============================================================================
# services/docx_service.py — Word Document (.docx) Generator
#
# Generates formatted Microsoft Word resume files for download.
# Called from candidates.py GET /candidates/{id}/download-resume.
#
# Two generation paths:
#
#   1. generate_word_from_structured(data: dict)
#      Used when the AI returned structured JSON (Groq/Gemini/Azure providers).
#      Data is already parsed into sections (name, skills, experience, etc.)
#      → Produces cleaner, better-formatted output.
#
#   2. generate_word_resume(resume_text: str)
#      Used as fallback for Ollama or older records without structured data.
#      Parses the plain text resume into sections using parse_resume_sections().
#      → Less precise but handles any text format.
#
# Output format: .docx binary bytes returned directly as an HTTP response.
# No file is written to disk — the document is generated in memory.
#
# Formatting: Navy/teal color scheme matching Kforce brand, Times New Roman,
# sections separated by teal divider lines (━━━━━━━).
# =============================================================================

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import io

DIVIDER = "━" * 35

NAVY  = RGBColor(0x1a, 0x1a, 0x2e)
TEAL  = RGBColor(0x2a, 0x9d, 0x8f)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK  = RGBColor(0x33, 0x33, 0x33)
GREY  = RGBColor(0x55, 0x55, 0x55)
LITE  = RGBColor(0x99, 0x99, 0x99)

BODY_FONT    = "Times New Roman"
DIVIDER_FONT = "MS Mincho"


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _p(doc, text="", bold=False, italic=False, size=10,
       color=DARK, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=1, space_after=3):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    para.alignment = align
    if text:
        run = para.add_run(text)
        run.bold           = bold
        run.italic         = italic
        run.font.size      = Pt(size)
        run.font.color.rgb = color
        run.font.name      = BODY_FONT
    return para


def _divider(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(DIVIDER)
    run.font.size      = Pt(8)
    run.font.color.rgb = TEAL
    run.font.name      = DIVIDER_FONT


def _section_header(doc, title: str):
    _divider(doc)
    _p(doc, title.upper(), bold=True, size=11, color=NAVY,
       space_before=6, space_after=2)
    _divider(doc)


def _bullet(doc, text: str):
    text = text.strip().lstrip("-•*· ").strip()
    if not text:
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.left_indent  = Inches(0.2)
    run = para.add_run(f"•  {text}")
    run.font.size      = Pt(10)
    run.font.color.rgb = DARK
    run.font.name      = BODY_FONT


def _company_line(doc, text: str):
    _p(doc, text, bold=True, size=10, color=NAVY,
       space_before=6, space_after=1)


def _project_line(doc, text: str):
    _p(doc, text, bold=True, size=10, color=BLACK,
       space_before=6, space_after=1)


def _sub_project_line(doc, text: str):
    _p(doc, text, bold=True, italic=True, size=10, color=BLACK,
       space_before=4, space_after=1)


# ── Junk-line filter (same patterns as llm_service.clean_resume_text) ─────────
import re as _re
_JUNK_LINE_PATTERNS = [
    r"^i am responsible for following",
    r"^i was responsible for following",
    r"^i am responsible for",
    r"^following are (the\s+)?responsibilities",
    r"^key responsibilities\s*:?$",
    r"^job profile\s*:",
    r"^roles\s*[&/and]*\s*responsibilities\s*:?$",
    r"^description\s*:?$",
    r"^technologies\s*[&/and]*\s*tools\s*:?$",
    r"^project description\s*:?$",
    r"^i hereby declare",
    r"^date of birth\s*:",
    r"^gender\s*:",
    r"^marital status\s*:",
    r"^date\s*:?\s*$",
    r"^personal details?\s*:?$",
    r"^declaration\s*:?$",
    r"^references?\s*(available)?\s*:?\s*$",
]

def _is_junk_line(line: str) -> bool:
    lower = line.strip().lower()
    return bool(lower) and any(_re.match(p, lower, _re.I) for p in _JUNK_LINE_PATTERNS)


# ── Skills auto-categorizer ───────────────────────────────────────────────────
_SKILL_CATS = [
    ("Automation Frameworks", {"selenium","playwright","cypress","testng","cucumber","pytest","jmeter","restassured","rest assured","appium","jasmine","mocha","jest","karate","gatling","locust","webdriver","detox"}),
    ("Languages",             {"java","python","javascript","typescript","sql","kotlin","scala","golang","ruby","php","csharp","swift","rust","bash","perl"}),
    ("AI / LLM Tools",        {"openai","chatgpt","claude","gemini","langchain","llamaindex","llm","rag","huggingface","pytorch","tensorflow","scikit","keras","mlflow","promptengineering","playwright agents","playwright agent"}),
    ("Cloud & DevOps",        {"aws","azure","gcp","kubernetes","docker","terraform","helm","jenkins","github actions","githubactions","circleci","gitlab","argocd","ansible"}),
    ("Databases",             {"mysql","postgresql","mongodb","redis","snowflake","elasticsearch","oracle","dynamodb","sqlite","cassandra","mssql","bigquery"}),
    ("Test Management",       {"jira","azure devops","azuredevops","testrail","zephyr","qtest","xray","confluence","bugzilla"}),
    ("API Testing",           {"postman","swagger","rest","graphql","soap","insomnia","karate"}),
    ("Methodologies",         {"agile","scrum","waterfall","kanban","bdd","tdd","devops","ci/cd","cicd"}),
    ("IDE & Tools",           {"vscode","vs code","eclipse","intellij","pycharm","cursor","git","github","bitbucket","maven","gradle"}),
]

def _categorize_skills(skills: list) -> dict:
    """Group a flat skills list into categories. Uncategorized go to 'Other'."""
    cats: dict = {}
    used: set = set()
    for cat_name, keywords in _SKILL_CATS:
        matched = []
        for skill in skills:
            sl = skill.lower().strip()
            if sl in used:
                continue
            if sl in keywords or any(kw in sl for kw in keywords):
                matched.append(skill)
                used.add(sl)
        if matched:
            cats[cat_name] = matched
    # Remaining
    other = [s for s in skills if s.lower().strip() not in used]
    if other:
        cats["Other Tools"] = other
    return cats


# ── Section parser ────────────────────────────────────────────────────────────

def parse_resume_sections(text: str) -> dict:
    sec = {
        "name": "", "role": "", "email": "", "mobile": "",
        "location": "", "linkedin": "",
        "summary": "", "skills": "", "experience": "",
        "projects": "", "education": "", "certifications": "",
        "raw": text,
    }

    # ── ━━━ structured format (from JD alignment) ─────────────────────────────
    if text.count("━") > 10:
        parts = [p.strip() for p in re.split(r"━{5,}", text) if p.strip()]

        if parts:
            hdr = [l.strip() for l in parts[0].split("\n") if l.strip()]
            if hdr:   sec["name"] = hdr[0]
            if len(hdr) > 1: sec["role"] = hdr[1]

        # contact
        for p in parts[:4]:
            for line in p.split("\n"):
                line = line.strip()
                m = re.search(r'[\w._%+\-]+@[\w.\-]+\.[a-z]{2,}', line, re.I)
                if m: sec["email"] = m.group()
                m = re.search(r'(\+91[\s\-]?)?[6-9]\d{9}', line)
                if m: sec["mobile"] = m.group()
                m = re.search(r'linkedin\.com\S*', line, re.I)
                if m: sec["linkedin"] = m.group()
                if re.search(r'📍|pune|mumbai|bangalore|hyderabad|chennai|delhi|noida', line, re.I):
                    clean = re.sub(r'[📍📧📱]', '', line).strip()
                    if len(clean) < 60: sec["location"] = clean

        mapping = {
            "PROFESSIONAL SUMMARY": "summary",
            "TECHNICAL SKILLS":     "skills",
            "WORK EXPERIENCE":      "experience",
            "PROJECT SUMMARY":      "projects",
            "EDUCATION":            "education",
            "CERTIFICATIONS":       "certifications",
        }
        for i, part in enumerate(parts):
            pu = part.strip().upper()
            for heading, key in mapping.items():
                if pu == heading:
                    if i + 1 < len(parts):
                        sec[key] = parts[i + 1]
                    break
        return sec

    # ── Raw PDF text fallback ─────────────────────────────────────────────────
    lines = text.split("\n")

    for line in lines[:15]:
        line = line.strip()
        if not line:
            continue
        m = re.search(r'[\w._%+\-]+@[\w.\-]+\.[a-z]{2,}', line, re.I)
        if m:
            sec["email"] = m.group(); continue
        m = re.search(r'(\+91[\s\-]?)?[6-9]\d{9}', line)
        if m:
            sec["mobile"] = m.group(); continue
        m = re.search(r'linkedin\.com\S*', line, re.I)
        if m:
            sec["linkedin"] = m.group(); continue
        if re.search(r'pune|mumbai|bangalore|hyderabad|chennai|delhi|noida|gurgaon|kolkata', line, re.I):
            sec["location"] = line; continue
        if not sec["name"] and re.match(r'^[A-Z][a-zA-Z ]{2,45}$', line):
            sec["name"] = line; continue
        if sec["name"] and not sec["role"] and len(line) < 80 and not re.search(r'@|\d{8,}', line):
            sec["role"] = line

    kw_map = {
        # Summary
        "summary": "summary", "professional summary": "summary",
        "objective": "summary", "profile": "summary", "about me": "summary",
        "career objective": "summary", "professional profile": "summary",
        "professional overview": "summary", "executive summary": "summary",
        # Skills
        "skills": "skills", "technical skills": "skills",
        "key skills": "skills", "core competencies": "skills",
        "expertise": "skills", "technologies": "skills",
        "skill set": "skills", "technical profic": "skills",
        "tools & technologies": "skills", "tools and technologies": "skills",
        "it skills": "skills", "computer skills": "skills",
        "technology stack": "skills", "tech stack": "skills",
        # Experience
        "experience": "experience", "work experience": "experience",
        "professional experience": "experience", "employment": "experience",
        "work history": "experience", "career history": "experience",
        "employment history": "experience", "professional background": "experience",
        # Projects
        "project": "projects", "project summary": "projects",
        "projects": "projects", "key projects": "projects",
        "notable projects": "projects", "portfolio": "projects",
        # Education
        "education": "education", "academic": "education",
        "qualification": "education", "educational": "education",
        "academic profile": "education", "academic background": "education",
        # Certifications
        "certif": "certifications", "certification": "certifications",
        "achievement": "certifications", "award": "certifications",
        "training": "certifications", "license": "certifications",
        # Personal details — map to ignore section
        "personal detail": "personal", "personal info": "personal",
        "personal information": "personal",
    }

    current = None
    buf: dict = {}
    for line in lines:
        s = line.strip()
        lower = s.lower()
        matched = None
        for kw, sname in kw_map.items():
            if lower.startswith(kw) and len(s) < 65:
                matched = sname
                break
        if matched:
            current = matched
            buf.setdefault(current, [])
        elif current and s:
            buf.setdefault(current, []).append(s)

    for k, v in buf.items():
        if k == "personal":
            continue  # drop personal details entirely
        sec[k] = "\n".join(v).strip()

    return sec


# ── Word document generator ───────────────────────────────────────────────────

def generate_word_resume(
    resume_text: str,
    candidate_name: str = "",
    version_label: str = "",
) -> bytes:

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    data = parse_resume_sections(resume_text)
    if candidate_name:
        data["name"] = candidate_name

    # ── HEADER ────────────────────────────────────────────────────────────────
    _divider(doc)

    np_ = doc.add_paragraph()
    np_.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    np_.paragraph_format.space_before = Pt(6)
    np_.paragraph_format.space_after  = Pt(2)
    nr = np_.add_run((data["name"] or "CANDIDATE NAME").upper())
    nr.bold = True
    nr.font.size      = Pt(20)
    nr.font.color.rgb = NAVY
    nr.font.name      = BODY_FONT

    if data["role"]:
        rp = doc.add_paragraph()
        rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rp.paragraph_format.space_after = Pt(2)
        rr = rp.add_run(data["role"])
        rr.bold           = True
        rr.font.size      = Pt(12)
        rr.font.color.rgb = BLACK
        rr.font.name      = BODY_FONT

    _divider(doc)

    # Contact line
    parts = []
    if data["email"]:    parts.append(f"📧 {data['email']}")
    if data["mobile"]:   parts.append(f"📱 {data['mobile']}")
    if data["location"]: parts.append(f"📍 {data['location']}")
    if parts:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cp.paragraph_format.space_after = Pt(2)
        cr = cp.add_run("  |  ".join(parts))
        cr.font.size      = Pt(9)
        cr.font.color.rgb = GREY
        cr.font.name      = BODY_FONT

    if data["linkedin"]:
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        lp.paragraph_format.space_after = Pt(4)
        lr = lp.add_run(f"LinkedIn: {data['linkedin']}")
        lr.font.size      = Pt(9)
        lr.font.color.rgb = TEAL
        lr.font.name      = BODY_FONT

    # Version label intentionally omitted — keep document clean for client submission

    # ── PROFESSIONAL SUMMARY ──────────────────────────────────────────────────
    if data["summary"]:
        _section_header(doc, "Professional Summary")
        for line in data["summary"].split("\n"):
            if line.strip():
                _p(doc, line.strip(), size=10, color=DARK)

    # ── TECHNICAL SKILLS ──────────────────────────────────────────────────────
    if data["skills"]:
        _section_header(doc, "Technical Skills")
        for line in data["skills"].split("\n"):
            line = line.strip()
            if not line:
                continue
            # Lines like "- Category: skill1, skill2"
            if re.match(r'^[-•]?\s*\w[\w\s]+:', line):
                _bullet(doc, line)
            else:
                _bullet(doc, line)

    # ── WORK EXPERIENCE ───────────────────────────────────────────────────────
    if data["experience"]:
        _section_header(doc, "Work Experience")
        prev_was_bullet = False
        for line in data["experience"].split("\n"):
            line = line.strip()
            if not line:
                prev_was_bullet = False
                continue
            # Strip junk boilerplate lines
            if _is_junk_line(line):
                continue
            # Company header: contains pipe OR has known company-identifying words
            if re.search(r'\|', line) and len(line) < 130:
                _company_line(doc, line)
                prev_was_bullet = False
            elif re.search(r'(?:Ltd|Pvt|Inc|Technologies|Solutions|Software|Corp|Systems|Limited|Birlasoft|Infosys|Wipro|TCS|Accenture|Cognizant|Capgemini|HCL|Tech Mahindra|Labs)\b', line, re.I) and len(line) < 120:
                _company_line(doc, line)
                prev_was_bullet = False
            elif re.match(r'^\d+\.\s+\w', line) and re.search(r'(?:Ltd|Pvt|Inc|Limited|Labs|Technologies)\b', line, re.I):
                # Numbered company entry like "1. Rysun Labs Pvt. Ltd..."
                _company_line(doc, re.sub(r'^\d+\.\s+', '', line))
                prev_was_bullet = False
            elif (re.search(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|20\d\d|19\d\d)', line, re.I)
                  and len(line) < 100
                  and not re.match(r'^[-•◦◆▪*]', line)
                  and not prev_was_bullet):
                # Line has a date and is short → sub-project heading (with date)
                _sub_project_line(doc, line)
                prev_was_bullet = False
            elif (re.search(r' [–\-] ', line)
                  and len(line) < 120
                  and not re.match(r'^[-•◦◆▪*]', line)
                  and re.match(r'^[A-Z]', line)
                  and not prev_was_bullet):
                # Title-case line with dash separator but no date → named sub-project heading
                # e.g. "Cement AI – Enterprise GenAI Platform (Dalmia Cement)"
                _sub_project_line(doc, line)
                prev_was_bullet = False
            elif re.match(r'^[-•◦◆▪*]', line) or re.match(r'^\w{1,4}[\.\)]\s', line):
                _bullet(doc, line)
                prev_was_bullet = True
            elif prev_was_bullet and re.match(r'^[a-z]', line):
                para = doc.paragraphs[-1]
                run = para.add_run(" " + line)
                run.font.size = Pt(10)
                run.font.color.rgb = DARK
            else:
                _bullet(doc, line)
                prev_was_bullet = True

    # ── PROJECT SUMMARY ───────────────────────────────────────────────────────
    if data["projects"]:
        _section_header(doc, "Project Summary")
        for line in data["projects"].split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("[DRAFT"):
                continue  # skip draft markers — clean output for client submission
            elif re.search(r'\|', line) and len(line) < 130:
                _project_line(doc, line)
            else:
                _bullet(doc, line)

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    if data["education"]:
        _section_header(doc, "Education")
        for line in data["education"].split("\n"):
            if line.strip():
                _p(doc, line.strip(), size=10, color=DARK)

    # ── CERTIFICATIONS ────────────────────────────────────────────────────────
    if data["certifications"]:
        _section_header(doc, "Certifications")
        for line in data["certifications"].split("\n"):
            if line.strip():
                _bullet(doc, line.strip())

    # ── FALLBACK: no sections found — dump full text ──────────────────────────
    no_content = not any([
        data["summary"], data["skills"], data["experience"],
        data["projects"], data["education"], data["certifications"],
    ])
    if no_content:
        _section_header(doc, "Resume Content")
        for line in resume_text.split("\n"):
            if line.strip():
                _p(doc, line.strip(), size=10, color=DARK)

    _divider(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── Structured JSON → Word doc ────────────────────────────────────────────────

def generate_word_from_structured(data: dict) -> bytes:
    """Generate a Word doc directly from structured JSON — no text parsing needed.
    Produces cleaner output than the text path since sections are already separated."""
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── HEADER ────────────────────────────────────────────────────────────────
    _divider(doc)

    np_ = doc.add_paragraph()
    np_.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    np_.paragraph_format.space_before = Pt(6)
    np_.paragraph_format.space_after  = Pt(2)
    nr = np_.add_run((data.get("name") or "CANDIDATE NAME").upper())
    nr.bold = True; nr.font.size = Pt(20); nr.font.color.rgb = NAVY; nr.font.name = BODY_FONT

    if data.get("role"):
        rp = doc.add_paragraph()
        rp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rp.paragraph_format.space_after = Pt(2)
        rr = rp.add_run(data["role"])
        rr.bold = True; rr.font.size = Pt(12); rr.font.color.rgb = BLACK; rr.font.name = BODY_FONT

    _divider(doc)

    contact = []
    if data.get("email"):    contact.append(f"📧 {data['email']}")
    if data.get("mobile"):   contact.append(f"📱 {data['mobile']}")
    if data.get("location"): contact.append(f"📍 {data['location']}")
    if contact:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        cp.paragraph_format.space_after = Pt(2)
        cr = cp.add_run("  |  ".join(contact))
        cr.font.size = Pt(9); cr.font.color.rgb = GREY; cr.font.name = BODY_FONT

    if data.get("linkedin"):
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        lp.paragraph_format.space_after = Pt(4)
        lr = lp.add_run(f"LinkedIn: {data['linkedin']}")
        lr.font.size = Pt(9); lr.font.color.rgb = TEAL; lr.font.name = BODY_FONT

    # ── PROFESSIONAL SUMMARY ──────────────────────────────────────────────────
    if data.get("summary"):
        _section_header(doc, "Professional Summary")
        _p(doc, data["summary"].strip(), size=10, color=DARK)

    # ── TECHNICAL SKILLS ──────────────────────────────────────────────────────
    skills = data.get("skills", [])
    if skills:
        _section_header(doc, "Technical Skills")
        if isinstance(skills, dict):
            # LLM returned categorized dict — render each category as a bullet line
            for cat_name, cat_items in skills.items():
                if cat_items:
                    items_str = ",  ".join(str(s) for s in cat_items)
                    _bullet(doc, f"{cat_name}:  {items_str}")
        elif isinstance(skills, list):
            cats = _categorize_skills(skills)
            if len(cats) > 1:
                for cat_name, cat_skills in cats.items():
                    _bullet(doc, f"{cat_name}:  {',  '.join(cat_skills)}")
            else:
                for i in range(0, len(skills), 6):
                    _bullet(doc, ",  ".join(str(s) for s in skills[i:i+6]))
        else:
            for line in str(skills).split("\n"):
                if line.strip():
                    _bullet(doc, line.strip())

    # ── WORK EXPERIENCE ───────────────────────────────────────────────────────
    # Structure: Company header → nested projects, each with period + tech + bullets.
    # Company-level bullets (rare) shown before the project list.
    experience = data.get("experience", [])
    if experience:
        _section_header(doc, "Work Experience")
        for job in experience:
            header_parts = [p for p in [job.get("company"), job.get("title"), job.get("period")] if p]
            if header_parts:
                _company_line(doc, "  |  ".join(header_parts))
            # Company-level bullets (if any — usually empty when projects are nested)
            for bullet in job.get("bullets", []):
                b = str(bullet).strip()
                if not b or _is_junk_line(b):
                    continue
                _bullet(doc, b)
            # Nested projects — each shown with name, period, tech, then bullets
            for proj in job.get("projects", []):
                proj_name   = str(proj.get("name", "")).strip()
                proj_period = str(proj.get("period", "")).strip()
                proj_tech   = str(proj.get("tech",   "")).strip()
                # Sub-project header: Name  |  Period
                name_period = "  |  ".join(p for p in [proj_name, proj_period] if p)
                if name_period:
                    _sub_project_line(doc, name_period)
                # Tech line (indented, smaller, grey)
                if proj_tech:
                    tp = doc.add_paragraph()
                    tp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    tp.paragraph_format.space_before = Pt(1)
                    tp.paragraph_format.space_after  = Pt(2)
                    tp.paragraph_format.left_indent  = Inches(0.25)
                    tr = tp.add_run(f"Tech: {proj_tech}")
                    tr.font.size      = Pt(9)
                    tr.font.italic    = True
                    tr.font.color.rgb = GREY
                    tr.font.name      = BODY_FONT
                # Project bullets
                for bullet in proj.get("bullets", []):
                    b = str(bullet).strip()
                    if b and not _is_junk_line(b):
                        _bullet(doc, b)

    # ── PROJECT SUMMARY (standalone/personal projects only) ──────────────────
    # Company-linked projects are shown nested under Work Experience above.
    standalone_projects = data.get("projects", [])
    if standalone_projects:
        _section_header(doc, "Project Summary")
        for proj in standalone_projects:
            proj_name   = str(proj.get("name", "")).strip()
            proj_period = str(proj.get("period", "")).strip()
            proj_tech   = str(proj.get("tech",   "")).strip()
            name_period = "  |  ".join(p for p in [proj_name, proj_period] if p)
            tech_str    = f"  |  Tech: {proj_tech}" if proj_tech else ""
            _project_line(doc, f"{name_period}{tech_str}")
            for bullet in proj.get("bullets", []):
                if str(bullet).strip():
                    _bullet(doc, str(bullet).strip())

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    education = data.get("education", [])
    if education:
        _section_header(doc, "Education")
        for edu in education:
            if isinstance(edu, dict):
                line = "  |  ".join(filter(None, [edu.get("degree"), edu.get("institution"), edu.get("year")]))
            else:
                line = str(edu)
            if line.strip():
                _p(doc, line.strip(), size=10, color=DARK)

    # ── CERTIFICATIONS ────────────────────────────────────────────────────────
    certifications = data.get("certifications", [])
    if certifications:
        _section_header(doc, "Certifications")
        for cert in certifications:
            if str(cert).strip():
                _bullet(doc, str(cert).strip())

    _divider(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
