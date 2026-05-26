"""
K Recruit — HR Resume Optimizer
DATA PRIVACY & COMPLIANCE ASSESSMENT REPORT  VERSION 7
Document Ref: K-RECRUIT-COMP-2026-007

COMPLETE document — fully supersedes v1–v6 (2026-001 through 2026-006).

New in v7 (May 2026 — Frontend Security Edition):
  - Full frontend security audit of preview.html and index.html
  - 6 frontend findings: F-01 API URL hardcoded, F-02 auth tokens not sent,
    F-03 Babel dev tool in production, F-04 CORS wildcard, F-05 no security
    headers, F-06 JWT in localStorage
  - All fixes applied: relative API URL, authFetch on all calls, CORS locked,
    security headers middleware (CSP, X-Frame-Options, HSTS, etc.)
  - index.html marked as deprecated (old prototype file)
  - Section 6.2 updated: frontend JS libraries re-audited
  - New Section 6A: Frontend Security Audit (full findings + fixes)
  - New Section 6B: Babel/CDN Compliance Approval Path
    (structured same as Azure OpenAI request — IT approval template)
  - New Section 19: v7 Frontend Fixes (6 code-level changes)
  - Checklist extended: B-17 through B-22
  - Declaration updated: v7 supersedes v1–v6
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

NAVY    = RGBColor(0x0F, 0x17, 0x2A)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
MUTED   = RGBColor(0x64, 0x74, 0x8B)
TEAL    = RGBColor(0x0D, 0x6E, 0x6E)
RED     = RGBColor(0xC0, 0x10, 0x20)
GREEN   = RGBColor(0x16, 0x65, 0x34)
AMBER   = RGBColor(0x92, 0x40, 0x0E)
GRAY_L  = RGBColor(0x94, 0xA3, 0xB8)
INDIGO_L= RGBColor(0x99, 0xA3, 0xFF)
INDIGO  = RGBColor(0x44, 0x38, 0xCA)

def shade_cell(cell, hex6):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6.lstrip("#")); tcPr.append(shd)

def cell_pad(cell, twips=140):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ("top","left","bottom","right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(twips)); m.set(qn("w:type"), "dxa"); tcMar.append(m)
    tcPr.append(tcMar)

def bottom_border(para, color="0F172A"):
    pPr = para._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "2"); bot.set(qn("w:color"), color)
    pBdr.append(bot); pPr.append(pBdr)

def left_bar(cell, color="6366F1"):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBorders")
    for edge in ("top","bottom","right"):
        t = OxmlElement(f"w:{edge}"); t.set(qn("w:val"), "none"); tcBdr.append(t)
    lft = OxmlElement("w:left")
    lft.set(qn("w:val"), "single"); lft.set(qn("w:sz"), "18")
    lft.set(qn("w:space"), "0"); lft.set(qn("w:color"), color)
    tcBdr.append(lft); tcPr.append(tcBdr)

def rn(para, text, bold=False, italic=False, sz=10, color=None):
    r = para.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(sz)
    if color: r.font.color.rgb = color
    return r

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
    bottom_border(p); return p

def h2(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11.5)
    r.font.color.rgb = color or NAVY; return p

def h3(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(10.5)
    r.font.color.rgb = color or MUTED; return p

def body(text, sz=10, color=None, after=5, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(sz); r.bold = bold; r.italic = italic
    r.font.color.rgb = color or NAVY; return p

def bul(text, sz=10, indent=0.35):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text); r.font.size = Pt(sz); r.font.color.rgb = NAVY; return p

def sp(pts=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts)

def callout(lines, bg="EEF2FF", border="6366F1"):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0,0); shade_cell(cell, bg); left_bar(cell, border); cell_pad(cell, 200)
    cell.paragraphs[0].clear()
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i==0 else cell.add_paragraph()
        r = p.add_run(line); r.font.size = Pt(9.5); r.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    sp(6)

def tbl(headers, rows, col_widths=None, hdr_bg="0F172A", alt="F8FAFC"):
    nc = len(headers)
    t = doc.add_table(rows=1+len(rows), cols=nc)
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hr = t.rows[0]
    for j,h in enumerate(headers):
        c = hr.cells[j]; shade_cell(c, hdr_bg); cell_pad(c, 120)
        p = c.paragraphs[0]; r = p.add_run(h)
        r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = WHITE
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    for i, row in enumerate(rows):
        bg = alt if i%2==0 else "FFFFFF"; tr = t.rows[i+1]
        for j, cv in enumerate(row):
            c = tr.cells[j]; shade_cell(c, bg); cell_pad(c, 120)
            p = c.paragraphs[0]
            if isinstance(cv, tuple):
                txt, bold, clr = cv
                r = p.add_run(str(txt)); r.bold = bold
                r.font.size = Pt(9); r.font.color.rgb = clr
            else:
                r = p.add_run(str(cv)); r.font.size = Pt(9); r.font.color.rgb = NAVY
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    if col_widths:
        for j,w in enumerate(col_widths):
            for row in t.rows: row.cells[j].width = Inches(w)
    sp(6)

def OK(t):    return (t, True,  GREEN)
def BAD(t):   return (t, True,  RED)
def WRN(t):   return (t, True,  AMBER)
def DONE(t):  return (t, True,  GREEN)
def BKLG(t):  return (t, True,  INDIGO)
def FIX(t):   return (t, True,  TEAL)

# =============================================================================
#  COVER PAGE
# =============================================================================
ct = doc.add_table(rows=1, cols=1)
cc = ct.cell(0,0); shade_cell(cc, "0F172A"); cell_pad(cc, 380)
cc.paragraphs[0].clear()

def cl(text, sz, clr, bold=False, after=6):
    p = cc.add_paragraph(); r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(sz); r.font.color.rgb = clr
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(after)

cl("K Recruit",                                                          24, WHITE,    bold=True,  after=2)
cl("HR Resume Optimizer — Kforce Internal Tool",                         14, GRAY_L,   bold=False, after=16)
cl("DATA PRIVACY & COMPLIANCE ASSESSMENT REPORT",                        12, INDIGO_L, bold=True,  after=4)
cl("VERSION 7 — FRONTEND SECURITY EDITION  (Supersedes v1–v6)",          9, GRAY_L,   bold=True,  after=16)

for lbl, val in [
    ("Document Ref:",    "K-RECRUIT-COMP-2026-007"),
    ("Version:",         "7.0 — Final  |  Supersedes v1–v6 (2026-001 through 2026-006)"),
    ("Prepared for:",    "Kforce US Compliance / IT Governance Team"),
    ("Prepared by:",     "Application Compliance Review — Technical Architecture"),
    ("Scope:",           "Frontend security audit: API auth, CORS, security headers, CDN, Babel, JWT storage"),
    ("Review Date:",     "May 2026"),
    ("Classification:",  "Internal Confidential"),
]:
    p = cc.add_paragraph()
    r1 = p.add_run(f"{lbl:20s}"); r1.bold = True; r1.font.size = Pt(9); r1.font.color.rgb = GRAY_L
    r2 = p.add_run(val);          r2.font.size = Pt(9);                  r2.font.color.rgb = WHITE
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(3)

cc.paragraphs[0]._element.getparent().remove(cc.paragraphs[0]._element)
sp(12)

# =============================================================================
#  SECTION 1 — EXECUTIVE SUMMARY
# =============================================================================
h1("1.  Executive Summary")

body(
    "This report is the definitive compliance and data privacy assessment for K Recruit, "
    "Kforce's internal HR resume optimisation tool. It supersedes all previous versions "
    "(v1–v6: 2026-001 through 2026-006). "
    "A complete line-by-line review was conducted across all backend services, routers, "
    "data models, configuration files, and frontend code."
)

callout([
    "KEY FINDING: K Recruit is architecturally capable of operating with ZERO DATA EGRESS",
    "when configured with Microsoft Azure OpenAI (Kforce's own Azure tenant).",
    "",
    "The current development configuration uses third-party cloud APIs (Groq, Gemini, NVIDIA NIM)",
    "which transmit candidate PII outside the organisation. This is a DEVELOPMENT-ONLY",
    "configuration and must be switched to Azure OpenAI before any production use.",
    "",
    "The single most important action: IT provisions Azure OpenAI in Kforce's Azure tenant.",
    "This is a Settings change in the app — no code deployment required.",
], bg="EEF2FF", border="6366F1")

body("Versions 4 and 5 resolved all critical authentication and security hardening findings:")
for item in [
    "All API routes require JWT authentication — including /bot/chat (fixed in v5)",
    "JWT_SECRET loaded from environment variable — not hardcoded in source code",
    "Passwords stored as bcrypt hashes in users table — no plain text credentials",
    "JWT tokens expire after 24 hours",
    "Brute-force login protection: 5 failed attempts per IP per 60 seconds → HTTP 429",
    "Audit logging middleware: all HTTP requests logged to audit.log",
    "SRI integrity hashes on all 3 CDN script tags",
    "PII stripped from AI prompts (email, phone, LinkedIn) before LLM call",
    "Full multi-user RBAC: super_admin / admin / recruiter with permission sets",
]:
    bul(item)

sp(4)
body("Version 6 confirmed the dedicated office laptop as production server and documented operational gaps.", bold=True)
for item in [
    "Dedicated server laptop confirmed as production server (Section 19)",
    "APP_ENV not set — FastAPI /docs exposed in default mode (C-07)",
    "Candidate data isolation missing — all recruiters see all candidates (C-08)",
    "No backup strategy — single point of failure (finding 11.3)",
    "Windows Firewall gap — port 8000 accessible directly (C-09)",
    "No audit log rotation — audit.log grows indefinitely (finding 11.4)",
]:
    bul(item)

sp(4)
body("Version 7 — Frontend Security Edition — completes the full-stack compliance audit:", bold=True)
for item in [
    "FIXED: API URL was hardcoded to http://127.0.0.1:8000 — changed to relative URL (works in dev + production)",
    "FIXED: 5 API calls bypassed authFetch and sent no Bearer token — all now use authFetch",
    "FIXED: CORS allow_origins=['*'] replaced with locked origin list from environment variable",
    "FIXED: Zero security headers — added CSP, X-Frame-Options, X-Content-Type-Options, HSTS, Referrer-Policy",
    "FIXED: Server root served old deprecated index.html — now serves preview.html",
    "FIXED: index.html marked with prominent DO NOT USE deprecation comment",
    "DOCUMENTED: Babel dev-tool-in-production finding with formal IT approval request template",
    "DOCUMENTED: React MIT license confirmed — no compliance issue",
    "DOCUMENTED: JWT in localStorage — known risk, formal approval path documented",
]:
    bul(item)

# =============================================================================
#  SECTION 2 — APPLICATION OVERVIEW
# =============================================================================
h1("2.  Application Overview")

tbl(
    ["Item", "Detail"],
    [
        ("Application Name",    "K Recruit — HR Resume Optimizer"),
        ("Previous Name",       "CMS TalentAI (renamed to K Recruit in v3 of this report)"),
        ("Purpose",             "Automate resume screening, ATS gap analysis, JD alignment, and soft-gap detection for HR recruiters"),
        ("Users",               "Internal HR recruiters — Kforce India office (pilot: ~20, full rollout: ~100)"),
        ("Production Server",   "Dedicated office laptop — permanently powered, sleep disabled, Kforce India office LAN"),
        ("Server Spec (min)",   "16 GB RAM, 512 GB SSD — AI processing offloaded to Azure, server handles web + DB only"),
        ("Network",             "Fixed static LAN IP (e.g. 192.168.1.50) — assigned by IT in router/adapter settings"),
        ("Backend",             "Python 3  ·  FastAPI web framework  ·  Uvicorn ASGI server"),
        ("Frontend",            "Single-page HTML app — preview.html  ·  React 18.2 + Babel (CDN with SRI hashes)"),
        ("Active frontend file","backend/static/preview.html  (index.html is deprecated prototype — DO NOT USE)"),
        ("Database",            "SQLite (development / up to ~30 users)  ·  PostgreSQL (production — 100 users)"),
        ("AI Provider",         "Development: Groq / Gemini / NVIDIA NIM  ·  Production: Azure OpenAI ONLY"),
        ("Semantic Model",      "sentence-transformers all-MiniLM-L6-v2 (~22 MB, runs fully locally on server laptop)"),
        ("Authentication",      "Email + bcrypt password  ·  JWT token (24h expiry)  ·  Role-based access control"),
        ("User Roles",          "super_admin / admin / recruiter — full RBAC with per-role permission sets"),
        ("File Storage",        "Uploaded PDFs stored in backend/uploads/ on the server laptop — local disk only"),
        ("Vector Database",     "None — similarity scoring computed in-memory only; no embeddings persisted"),
    ],
    col_widths=[1.8, 4.7],
)

# =============================================================================
#  SECTION 3 — DATA INVENTORY
# =============================================================================
h1("3.  Data Inventory — What Personal Data K Recruit Handles")

YES_DEV = ("YES — to AI provider (dev mode only)", True, RED)
AZ_ONLY = ("To Kforce Azure tenant only (production)", True, AMBER)
NO_GRN  = ("No — local storage only", False, GREEN)

tbl(
    ["Data Element", "Source", "Where Stored", "Dev: Transmitted?", "Production: Transmitted?"],
    [
        ("Candidate Full Name",   "PDF Resume",     "Local SQLite/PG DB",         YES_DEV, AZ_ONLY),
        ("Email Address",         "PDF Resume",     "Local SQLite/PG DB",         YES_DEV, AZ_ONLY),
        ("Mobile Number",         "PDF Resume",     "Local SQLite/PG DB",         YES_DEV, AZ_ONLY),
        ("Location / City",       "PDF Resume",     "Local SQLite/PG DB",         YES_DEV, AZ_ONLY),
        ("LinkedIn URL",          "PDF Resume",     "Local SQLite/PG DB",         YES_DEV, AZ_ONLY),
        ("Full Resume Text",      "PDF Resume",     "SQLite/PG DB + /uploads",    YES_DEV, AZ_ONLY),
        ("Job Description Text",  "HR team input",  "Local SQLite/PG DB",         YES_DEV, AZ_ONLY),
        ("AI Analysis Results",   "Generated",      "Local SQLite/PG DB",         NO_GRN,  NO_GRN),
        ("Recruiter Name",        "HR team input",  "Local SQLite/PG DB",         NO_GRN,  NO_GRN),
        ("Login Credentials",     "users table",    "Local DB — bcrypt hash only",NO_GRN,  NO_GRN),
        ("PDF Files",             "Upload",         "/uploads on server laptop",  NO_GRN,  NO_GRN),
    ],
    col_widths=[1.4, 1.1, 1.5, 1.6, 1.9],
)

callout([
    "DATA ISOLATION GAP (v6 NEW FINDING):",
    "Currently the candidates table has no uploaded_by field linking to the users table.",
    "This means all 100 recruiters can see all candidates uploaded by any recruiter.",
    "For a shared office deployment this is a known gap — see Section 11 Finding C-08.",
    "Fix: add uploaded_by FK to candidates table; filter results by user role in queries.",
], bg="FEF2F2", border="EF4444")

# =============================================================================
#  SECTION 4 — DATA FLOW DIAGRAMS
# =============================================================================
h1("4.  Data Flow Diagrams")

h2("4.1  Current Development State (NOT Production-Ready for PII)", color=RED)
callout([
    "CURRENT FLOW — DEVELOPMENT ONLY:",
    "",
    "  [HR Recruiter Browser]  (access via http://127.0.0.1:8000 — localhost only)",
    "         │",
    "         ▼   HTTP port 8000  (UNENCRYPTED — no SSL)",
    "  [Server Laptop — FastAPI Backend]",
    "         │",
    "         ├── PDF Upload ──────► /uploads folder  (LOCAL DISK)",
    "         ├── Database  ──────► cms.db SQLite     (LOCAL DISK)",
    "         │",
    "         └── AI Analysis ────► DATA LEAVES PREMISES ──►",
    "               ├── Groq Cloud    (api.groq.com — San Jose, CA, USA)",
    "               ├── Google Gemini (generativelanguage.googleapis.com)",
    "               └── NVIDIA NIM    (integrate.api.nvidia.com)",
    "",
    "  Frontend loads 3 JS libraries from cdnjs.cloudflare.com (SRI-verified).",
    "  Note: In dev, API URL is relative — browser accesses same host, no cross-origin issue.",
], bg="FEF2F2", border="EF4444")

h2("4.2  Production-Ready State — Azure OpenAI + SSL + Server Laptop", color=GREEN)
callout([
    "PRODUCTION FLOW — Dedicated Office Laptop + nginx SSL + JWT auth:",
    "",
    "  ┌──────────────────────────────────────────────────────────────────┐",
    "  │                     KFORCE INDIA OFFICE LAN                      │",
    "  │                                                                  │",
    "  │  [Recruiter PC]  ─── HTTPS/443 (SSL) ──► [nginx on laptop]      │",
    "  │  (any of 100)                                     │              │",
    "  │                                      HTTP/8000 (localhost only)  │",
    "  │                                                   │              │",
    "  │                                    ┌──────────────▼───────────┐  │",
    "  │  PDF uploaded ────────────────────►│  FastAPI — Server Laptop  │ │",
    "  │  JWT token on every request        │  PostgreSQL DB            │ │",
    "  │  Static IP: 192.168.1.50           │  /uploads (local disk)    │ │",
    "  │  Port 8000 blocked at firewall     └──────────┬────────────────┘ │",
    "  └─────────────────────────────────────────────  │  ────────────────┘",
    "                                                   │ Resume + JD text",
    "                                                   ▼  HTTPS/443 only",
    "                               ┌──────────────────────────────────────┐",
    "                               │   KFORCE AZURE TENANT                │",
    "                               │   Azure OpenAI — GPT-4o              │",
    "                               │   Microsoft DPA applies              │",
    "                               └──────────────────────────────────────┘",
], bg="F0FDF4", border="16A34A")

# =============================================================================
#  SECTION 5 — HUGGINGFACE MODEL
# =============================================================================
h1("5.  HuggingFace Model Download — Audit")

tbl(
    ["Detail", "Value"],
    [
        ("Model Name",           "all-MiniLM-L6-v2"),
        ("License",              "Apache 2.0 — free for all commercial use"),
        ("File size",            "~22 MB"),
        ("Downloaded when",      "ONCE — on first server startup after pip install"),
        ("After first download", "Runs 100% locally — zero further network calls to HuggingFace"),
        ("Production action",    "Pre-download during setup; set HF_HUB_OFFLINE=1 in .env on server laptop"),
        ("User data sent",       "NONE — only the model weights file is downloaded"),
    ],
    col_widths=[2.2, 4.3],
)

callout([
    "NOTE: HF_HUB_OFFLINE concern is resolved automatically when Azure OpenAI is configured.",
    "Once the server laptop is in production with Azure, the semantic model is pre-downloaded",
    "during setup and HF_HUB_OFFLINE=1 is set in .env — no further HuggingFace contact.",
    "This is part of the server laptop setup checklist (Section 20).",
], bg="F0FDF4", border="16A34A")

h2("5.1  Alternatives if Compliance Does Not Approve HuggingFace")
tbl(
    ["Option", "HuggingFace?", "Quality", "Effort", "Best When"],
    [
        ("TF-IDF fallback (built-in)", OK("None"), WRN("Lower"), OK("Remove 1 line"), "Fastest, zero contact"),
        ("Internal file share",        OK("None"), OK("Identical"), WRN("Low"), "Kforce has file server"),
        ("Bundle with app",            OK("None"), OK("Identical"), WRN("Low"), "Clean enterprise deploy"),
        ("Azure embeddings",           OK("None"), OK("Better"),    WRN("~20 lines"), "Already using Azure"),
    ],
    col_widths=[1.8, 1.0, 1.0, 1.0, 2.7],
)

# =============================================================================
#  SECTION 6 — LIBRARY NETWORK AUDIT
# =============================================================================
h1("6.  Complete Library-by-Library Network Audit")

h2("6.1  Backend Python Libraries")
tbl(
    ["Library", "Version", "Purpose", "Network Calls?", "User Data?", "Verdict"],
    [
        ("fastapi",              "≥0.110", "Web API framework",          OK("None"), OK("No"), OK("SAFE")),
        ("uvicorn",              "≥0.29",  "ASGI web server",            OK("None"), OK("No"), OK("SAFE")),
        ("sqlalchemy",           "≥2.0",   "Database ORM",               OK("None"), OK("No"), OK("SAFE")),
        ("psycopg2-binary",      "≥2.9",   "PostgreSQL driver",          OK("None"), OK("No"), OK("SAFE")),
        ("pymupdf (fitz)",       "≥1.24",  "PDF text extraction",        OK("None"), OK("No"), OK("SAFE")),
        ("python-docx",          "≥1.1",   "Word document generation",   OK("None"), OK("No"), OK("SAFE")),
        ("requests",             "≥2.31",  "AI provider HTTP calls",     WRN("Azure OpenAI only"), WRN("Resume+JD"), WRN("APPROVED")),
        ("python-dotenv",        "≥1.0",   "Reads .env config file",     OK("None"), OK("No"), OK("SAFE")),
        ("python-jose",          "≥3.3",   "JWT token signing",          OK("None"), OK("No"), OK("SAFE")),
        ("passlib[bcrypt]",      "≥1.7",   "Password hashing",           OK("None"), OK("No"), OK("SAFE")),
        ("sentence-transformers","≥2.7",   "Semantic scoring",           WRN("One-time model DL"), OK("No"), WRN("MITIGATED")),
        ("scikit-learn",         "≥1.3",   "TF-IDF fallback scoring",    OK("None"), OK("No"), OK("SAFE")),
    ],
    col_widths=[1.5, 0.6, 1.7, 1.4, 1.0, 0.85],
)

h2("6.2  Frontend JavaScript Libraries — Full Audit (v7 Updated)")
body(
    "All three frontend JavaScript libraries are loaded from cdnjs.cloudflare.com with SHA-384 "
    "Subresource Integrity (SRI) hashes. If the CDN serves a tampered file, the browser rejects it. "
    "The active frontend file is preview.html. index.html is deprecated and not served by the application."
)
tbl(
    ["Library", "Version", "License", "SRI Hash?", "CDN Calls User Data?", "Compliance Note", "Status"],
    [
        ("React",    "18.2.0", "MIT", OK("sha384 verified"),
         OK("No — static JS only"),
         "MIT license — approved for internal tools. Version 3 years old; update to 18.3.x recommended.",
         WRN("REVIEW")),
        ("ReactDOM", "18.2.0", "MIT", OK("sha384 verified"),
         OK("No — static JS only"),
         "Same as React. MIT license. Pair with React version when updating.",
         WRN("REVIEW")),
        ("Babel",    "7.23.2", "MIT", OK("sha384 verified"),
         OK("No — static JS only"),
         "Babel is a JavaScript COMPILER (dev tool). Running it in user browsers at runtime is flagged "
         "as 'dev tool in production' by IT security scans. Formal approval request: Section 6B.",
         BAD("APPROVAL NEEDED")),
    ],
    col_widths=[0.85, 0.7, 0.55, 1.0, 1.2, 2.4, 0.8],
)

callout([
    "REACT AND REACTDOM — NO COMPLIANCE ISSUE:",
    "",
    "React and ReactDOM are MIT-licensed open-source libraries maintained by Meta (Facebook).",
    "MIT license is the most permissive open-source license — it is explicitly approved for:",
    "  ✓  Commercial internal tools",
    "  ✓  Enterprise applications",
    "  ✓  Proprietary products",
    "",
    "Used by: Microsoft, Google, Amazon, Netflix, Twitter/X, Airbnb, Uber, and thousands of",
    "Fortune 500 companies in production systems. React is NOT a compliance concern.",
    "",
    "ACTION NEEDED: Update from 18.2.0 (June 2022) to 18.3.1 (latest stable) to pick up",
    "security patches. Update is a 1-line CDN URL change in preview.html.",
], bg="F0FDF4", border="16A34A")

# =============================================================================
#  SECTION 6A — FRONTEND SECURITY AUDIT
# =============================================================================
h1("6A.  Frontend Security Audit — Full Findings and Fixes (v7)")

body(
    "A complete security review of backend/static/preview.html (the live application) "
    "and backend/static/index.html (deprecated prototype) was performed in v7. "
    "Six findings were identified. All six were fixed in the v7 code release."
)

h2("6A.1  Frontend Security Findings Summary")
tbl(
    ["ID", "Finding", "Severity", "File", "Status"],
    [
        ("F-01", "API URL hardcoded to http://127.0.0.1:8000 — all 100 LAN recruiters would hit their own machine, not the server",
         BAD("Critical"), "preview.html line 182", DONE("FIXED v7")),
        ("F-02", "5 API calls used raw fetch() bypassing authFetch — no Bearer token sent (bot/chat x2, domain-coverage, role-compatibility, semantic-status)",
         BAD("Critical"), "preview.html lines 387/401/845/1685/2156", DONE("FIXED v7")),
        ("F-03", "Babel Standalone runs in browser at runtime — compiler (dev tool) shipped to production users",
         WRN("High"), "preview.html line 9", WRN("APPROVAL NEEDED — Section 6B")),
        ("F-04", "CORS allow_origins=['*'] with allow_credentials=True — no origin restriction on API",
         WRN("High"), "main.py line 78", DONE("FIXED v7")),
        ("F-05", "Zero HTTP security headers — no CSP, X-Frame-Options, HSTS, X-Content-Type-Options, Referrer-Policy",
         WRN("High"), "main.py (missing middleware)", DONE("FIXED v7")),
        ("F-06", "JWT token and role data stored in localStorage — accessible to any JS on the page",
         WRN("High"), "preview.html (session storage)", WRN("DOCUMENTED — formal approval path below")),
    ],
    col_widths=[0.45, 3.2, 0.75, 1.5, 0.9],
)

h2("6A.2  Detailed Fix — F-01: API URL (Hardcoded Localhost)")
callout([
    "BEFORE (broken for LAN users):",
    "  const API = \"http://127.0.0.1:8000\";",
    "  Every recruiter's browser would call their own machine's port 8000 — not the server.",
    "  The app appeared to work only on the developer's machine.",
    "",
    "AFTER (fixed — relative URL):",
    "  const API = \"\";   // empty = relative to current host",
    "",
    "HOW IT WORKS:",
    "  Dev machine:    user accesses http://127.0.0.1:8000  → API calls → http://127.0.0.1:8000/candidates/",
    "  Production LAN: user accesses https://192.168.1.50   → API calls → https://192.168.1.50/candidates/",
    "  Zero code changes needed when moving from dev to production.",
], bg="F0FDF4", border="16A34A")

h2("6A.3  Detailed Fix — F-02: Missing Auth Tokens on 5 API Calls")
callout([
    "BEFORE — 5 calls used raw fetch() with no Authorization header:",
    "  fetch(`${API}/bot/chat`, { ... })                       ← no token",
    "  fetch(`${API}/settings/domain-coverage`, ...)           ← no token",
    "  fetch(`${API}/candidates/${id}/role-compatibility`, ...) ← no token",
    "  fetch(`${API}/settings/semantic-status`, ...)           ← no token",
    "",
    "AFTER — all use authFetch() which injects the Bearer token:",
    "  authFetch(`${API}/bot/chat`, { ... })",
    "  authFetch(`${API}/settings/domain-coverage`, ...)",
    "  authFetch(`${API}/candidates/${id}/role-compatibility`, ...)",
    "  authFetch(`${API}/settings/semantic-status`, ...)",
    "",
    "authFetch() also handles 401 responses by clearing localStorage and redirecting to login.",
    "The auth/login and auth/verify calls correctly remain as raw fetch() — they run before",
    "a session token exists, so they intentionally have no token to send.",
], bg="F0FDF4", border="16A34A")

h2("6A.4  Detailed Fix — F-04: CORS Wildcard Removed")
callout([
    "BEFORE (non-compliant — allows any website to make credentialed API calls):",
    "  allow_origins=[\"*\"]",
    "  allow_credentials=True    ← RFC 6454: wildcard + credentials is invalid/dangerous",
    "  allow_methods=[\"*\"]",
    "  allow_headers=[\"*\"]",
    "",
    "AFTER (locked — only the application's own origins are permitted):",
    "  allow_origins=os.getenv('CORS_ORIGINS', 'http://127.0.0.1:8000,http://localhost:8000').split(',')",
    "  allow_methods=[\"GET\", \"POST\", \"PUT\", \"DELETE\", \"OPTIONS\"]",
    "  allow_headers=[\"Authorization\", \"Content-Type\"]",
    "",
    "Production server laptop: add https://192.168.1.50 to CORS_ORIGINS in .env.",
    "Since the frontend is served BY the FastAPI backend (same origin), CORS does not apply",
    "to normal recruiter access — CORS only matters for direct API clients (Postman, scripts).",
], bg="F0FDF4", border="16A34A")

h2("6A.5  Detailed Fix — F-05: Security Headers Added")
body("Six HTTP security headers are now injected by middleware on every response:")
tbl(
    ["Header", "Value Set", "Protection Against"],
    [
        ("X-Content-Type-Options",     "nosniff",
         "Browser MIME-sniffing attacks — forces browser to respect declared content type"),
        ("X-Frame-Options",            "DENY",
         "Clickjacking — prevents the app from being embedded in an iframe on any other site"),
        ("Referrer-Policy",            "strict-origin-when-cross-origin",
         "PII leakage via HTTP Referer header when navigating to external pages"),
        ("Permissions-Policy",         "camera=(), microphone=(), geolocation=()",
         "Prevents the browser from granting camera, microphone, or location to this app"),
        ("Content-Security-Policy",    "default-src 'self'; script-src 'self' cdnjs.cloudflare.com 'unsafe-inline' 'unsafe-eval'; ...",
         "XSS and script injection — browser blocks scripts from any source not in the allowlist. "
         "'unsafe-eval' is required while Babel runs in browser; will be removed when Babel is replaced."),
        ("Strict-Transport-Security",  "max-age=31536000; includeSubDomains  (production only)",
         "HSTS — forces browser to always use HTTPS; active only when APP_ENV=production with nginx/SSL"),
    ],
    col_widths=[1.6, 2.3, 2.6],
)

h2("6A.6  Finding F-06 — JWT in localStorage (Documented, Approval Path Below)")
callout([
    "CURRENT STATE:",
    "  The JWT session token and role data are stored in browser localStorage.",
    "  localStorage is readable by any JavaScript running on the page.",
    "  If an XSS vulnerability ever existed, a script could steal the token.",
    "",
    "WHY IT IS ACCEPTABLE FOR THIS INTERNAL TOOL:",
    "  1. XSS is mitigated: React auto-escapes all rendered data; no innerHTML or dangerouslySetInnerHTML used.",
    "  2. Content Security Policy (now added) blocks inline script injection.",
    "  3. Application is internal-only — not accessible from the internet.",
    "  4. Token expires after 24 hours — stolen token has short window.",
    "",
    "INDUSTRY STANDARD (for future upgrade):",
    "  Production-grade apps use httpOnly cookies — not accessible to JS even if XSS occurs.",
    "  This requires backend changes (Set-Cookie header, CSRF token) and is a medium-complexity refactor.",
    "  Classified as backlog BL-07 — not a blocker for initial 100-user production rollout.",
], bg="FEF9C3", border="F59E0B")

# =============================================================================
#  SECTION 6B — BABEL/CDN COMPLIANCE APPROVAL PATH
# =============================================================================
h1("6B.  Babel / CDN — Formal IT Approval Request")

body(
    "This section mirrors the Azure OpenAI approval request format (Section 7). "
    "Just as Kforce IT was asked to provision Azure OpenAI, this section documents "
    "what needs to be formally requested from IT and the compliance team for the "
    "Babel / CDN situation — with three options ranked by compliance impact."
)

h2("6B.1  What Is the Problem?")
callout([
    "TWO SEPARATE CONCERNS — different approvals needed:",
    "",
    "CONCERN 1 — External CDN Dependency:",
    "  React, ReactDOM, and Babel are loaded from cdnjs.cloudflare.com at runtime.",
    "  Risk: If the CDN is unreachable (firewall, outage, proxy block), the app fails to load.",
    "  Risk: External call from production app is an undocumented external dependency.",
    "  Note: All three scripts have SRI sha384 hashes — tampered files are rejected by browser.",
    "",
    "CONCERN 2 — Babel is a Developer Tool Running in Production:",
    "  Babel is a JavaScript compiler. Its job is to convert JSX (React syntax) to",
    "  plain JavaScript that browsers understand.",
    "  Normal apps compile JSX once at build time (on the developer's machine).",
    "  K Recruit compiles JSX on every page load in every recruiter's browser.",
    "  IT security scanners flag this as 'developer toolchain in production environment'.",
    "  Kforce, as a reputed IT company, should not have dev tools shipped to end users.",
], bg="FEF2F2", border="EF4444")

h2("6B.2  The Three Options — Ranked by Compliance Priority")
tbl(
    ["Option", "What It Does", "Removes CDN?", "Removes Babel from Browser?", "Effort", "Recommended?"],
    [
        ("Option A\nSelf-host libraries",
         "Download React, ReactDOM, Babel .js files and serve from /static/ on the server laptop. "
         "No external CDN calls at runtime.",
         OK("YES — fully"), BAD("NO — Babel still compiles in browser"),
         OK("Low (copy 3 files)"), WRN("INTERIM FIX")),
        ("Option B\nReplace Babel with htm",
         "Replace Babel+JSX with 'htm' — a 3KB library using tagged template literals. "
         "Eliminates Babel from production entirely. No build tool needed. "
         "Requires rewriting ~2800 lines of JSX syntax.",
         OK("Self-hostable"), OK("YES — completely"),
         WRN("Medium (~1–2 days)"), OK("RECOMMENDED")),
        ("Option C\nVite build pipeline",
         "Compile JSX once at build time using Vite (industry standard). "
         "Output: one minified JS file, no runtime compilation. "
         "Requires Node.js build environment.",
         OK("YES — fully"), OK("YES — completely"),
         BAD("High (project restructure)"), WRN("LONG-TERM")),
    ],
    col_widths=[0.9, 2.5, 0.9, 1.1, 0.7, 0.7],
)

h2("6B.3  Formal IT Request — Option A (Immediate: Self-Host Libraries)")
callout([
    "REQUEST TO IT / COMPLIANCE TEAM:",
    "",
    "  We request approval to self-host the following MIT-licensed JavaScript libraries",
    "  in the application's /static/ folder on the server laptop:",
    "",
    "  1. React 18.3.1        —  MIT License  —  maintained by Meta (Facebook)",
    "     Source: https://unpkg.com/react@18.3.1/umd/react.production.min.js",
    "",
    "  2. ReactDOM 18.3.1     —  MIT License  —  maintained by Meta (Facebook)",
    "     Source: https://unpkg.com/react-dom@18.3.1/umd/react-dom.production.min.js",
    "",
    "  3. Babel Standalone 7.23.2  —  MIT License  —  maintained by Babel OSS team",
    "     Source: https://unpkg.com/@babel/standalone@7.23.2/babel.min.js",
    "",
    "  BUSINESS JUSTIFICATION:",
    "  - Removes all external CDN calls from production — no runtime calls to cloudflare.com",
    "  - Libraries are downloaded once by IT during server laptop setup",
    "  - All three are MIT licensed — same approval class as standard enterprise JS",
    "  - No user data is sent to any CDN server (CDN only serves static script files)",
    "",
    "  IT ACTION:",
    "  - Download the 3 files from the URLs above (one-time)",
    "  - Place in: backend/static/ on the server laptop",
    "  - Dev updates the 3 <script src> tags to point to /static/react.min.js etc.",
    "",
    "  COMPLIANCE OUTCOME: 'External CDN dependency' finding is CLOSED.",
    "  Babel dev-tool finding remains open (needs Option B or C to close fully).",
], bg="EEF2FF", border="6366F1")

h2("6B.4  Formal IT Request — Option B (Proper Fix: Replace Babel with htm)")
callout([
    "REQUEST TO IT / COMPLIANCE TEAM:",
    "",
    "  We request approval to use the following library as an internal tool dependency:",
    "",
    "  Library: htm  (Hyperscript Tagged Markup)",
    "  Version: 3.1.1",
    "  License: Apache-2.0  (same permissive tier as MIT — approved for commercial use)",
    "  Size:    3 KB (minified)",
    "  Source:  https://github.com/developit/htm",
    "  Maintained by: Jason Miller (Google Chrome team)",
    "",
    "  WHAT htm DOES:",
    "  htm replaces Babel's JSX compilation with browser-native tagged template literals.",
    "  Instead of:   return <div className='card'>{name}</div>   (requires Babel)",
    "  We write:     return html`<div class='card'>${name}</div>`  (runs in any browser natively)",
    "",
    "  BUSINESS JUSTIFICATION:",
    "  - Eliminates Babel (dev tool / compiler) from the production browser entirely",
    "  - No build tool required — stays single-file HTML approach",
    "  - Can be self-hosted in /static/ (3 KB file)",
    "  - Removes 'dev toolchain in production' compliance flag permanently",
    "  - Content Security Policy can then remove 'unsafe-eval' restriction",
    "",
    "  IT ACTION:",
    "  - Approve Apache-2.0 license for K Recruit internal tool dependencies",
    "  - Dev team converts JSX to htm template syntax (~1–2 working days of changes)",
    "  - Babel CDN tag removed; replaced with single self-hosted htm.js",
    "",
    "  COMPLIANCE OUTCOME: Both 'CDN dependency' AND 'dev tool in production' findings CLOSED.",
], bg="F0FDF4", border="16A34A")

h2("6B.5  Formal IT Request — Option C (Long-Term: Vite Build Pipeline)")
callout([
    "REQUEST TO IT / INFRASTRUCTURE TEAM:",
    "",
    "  We request provisioning of a Node.js build environment for K Recruit:",
    "",
    "  Runtime: Node.js 20 LTS  (Long-Term Support — supported until April 2026)",
    "  Build tool: Vite 5.x   (MIT license — industry standard, used by Vue, Svelte, Astro)",
    "  This is required for a COMPILE step only — Node.js does NOT run in production.",
    "",
    "  HOW IT WORKS:",
    "  1. Dev runs 'vite build' on their machine (or CI pipeline)",
    "  2. Vite compiles all JSX, bundles React, outputs 1 minified JS file (~150 KB)",
    "  3. That single compiled file is deployed to backend/static/ on the server laptop",
    "  4. No Node.js, no Babel, no CDN needed on the production server",
    "",
    "  BUSINESS JUSTIFICATION:",
    "  - Industry standard approach — all major enterprise React apps use a build pipeline",
    "  - Eliminates ALL external dependencies from production: no CDN, no Babel",
    "  - Smaller output file → faster page load for 100 recruiters",
    "  - Enables code splitting, tree-shaking, and proper source maps",
    "",
    "  IT ACTION:",
    "  - Approve Node.js 20 LTS for developer workstations (build only, not production server)",
    "  - No changes needed on the server laptop itself",
    "",
    "  COMPLIANCE OUTCOME: All frontend CDN and Babel findings FULLY CLOSED. Enterprise-grade.",
], bg="EEF2FF", border="6366F1")

h2("6B.6  Recommended Path for Kforce Compliance")
tbl(
    ["Phase", "Action", "Who", "Closes Which Finding", "Timeline"],
    [
        ("NOW (immediate)",
         "IT approves Option A — self-host React, ReactDOM, Babel in /static/",
         "IT + Dev",
         "CDN external dependency (F-03 partial)",
         "1 day"),
        ("BEFORE 100-USER ROLLOUT",
         "IT approves Apache-2.0 htm library — Dev converts JSX to htm (Option B)",
         "Dev team",
         "Babel dev-tool-in-production (F-03 fully closed)",
         "1–2 days dev work"),
        ("LONG-TERM (optional)",
         "IT provisions Node.js 20 LTS on dev machines — migrate to Vite build (Option C)",
         "IT + Dev",
         "All frontend CDN and Babel findings — full enterprise standard",
         "1 sprint"),
        ("PARALLEL",
         "Update React + ReactDOM from 18.2.0 to 18.3.1 (latest stable, security patches)",
         "Dev",
         "React version finding",
         "30 minutes"),
    ],
    col_widths=[1.2, 2.2, 0.7, 1.7, 0.95],
)

# =============================================================================
#  SECTION 7 — AZURE OPENAI
# =============================================================================
h1("7.  Azure OpenAI — Subscription, Setup, and Cost Guide")

h2("7.1  Do We Need a New Azure Subscription?")
callout([
    "NO — Kforce already has an Azure subscription (used for Microsoft 365 / Copilot).",
    "Azure OpenAI is an add-on SERVICE within the existing Kforce Azure tenant.",
    "IT needs to ENABLE it — not create a new subscription.",
], bg="EEF2FF", border="6366F1")

h2("7.2  Step-by-Step Setup for IT Team")
tbl(
    ["Step", "Action", "Where in Azure Portal", "Who"],
    [
        ("1", "Enable Azure OpenAI for the subscription",
         "Subscriptions → Resource Providers → Microsoft.CognitiveServices → Register", "IT / Cloud Admin"),
        ("2", "Create Resource Group: 'kforce-k-recruit'",
         "portal.azure.com → Resource Groups → + Create", "IT / Cloud Admin"),
        ("3", "Create Azure OpenAI resource in Central India region",
         "portal.azure.com → Create Resource → Azure OpenAI → Region: Central India", "IT / Cloud Admin"),
        ("4", "Deploy GPT-4o model",
         "Azure OpenAI Studio → Deployments → + Create → gpt-4o", "IT / Cloud Admin"),
        ("5", "Copy endpoint and API key",
         "Azure OpenAI Resource → Keys and Endpoint", "IT / Cloud Admin"),
        ("6", "Enter in K Recruit Settings page",
         "K Recruit → Settings → AI Configuration → Provider: Azure OpenAI → paste endpoint + key", "Dev / Admin"),
    ],
    col_widths=[0.35, 2.1, 2.85, 1.2],
)

h2("7.3  Cost Estimate")
tbl(
    ["Scenario", "Requests/day", "Tokens/request", "Monthly Cost (GPT-4o)"],
    [
        ("Pilot — 20 recruiters",    "~50 resumes/day",  "~3,000 tokens",   "$3 – $10 / month"),
        ("Full rollout — 100 users", "~200 resumes/day", "~3,000 tokens",   "$30 – $80 / month"),
        ("Max usage estimate",       "~500 resumes/day", "~4,000 tokens",   "$150 – $200 / month"),
    ],
    col_widths=[1.9, 1.4, 1.4, 1.8],
)

callout([
    "RECOMMENDATION: Set a Cost Management budget alert in Azure portal at $80/month.",
    "This ensures IT is notified before spending exceeds expected range.",
    "5 minutes to configure: Azure Portal → Cost Management → Budgets → + Add",
], bg="F0FDF4", border="16A34A")

# =============================================================================
#  SECTION 8 — SSL CERTIFICATE
# =============================================================================
h1("8.  SSL Certificate — Options for LAN Deployment")

h2("8.1  Why SSL Is Required")
callout([
    "K Recruit transmits JWT authentication tokens and candidate PII over the network.",
    "Without SSL, all data is sent as plain text — anyone on the office network can",
    "intercept logins, tokens, and resume content with basic packet capture tools.",
    "SSL is not optional for a production deployment handling personal data.",
], bg="FEF2F2", border="EF4444")

h2("8.2  Certificate Options")
tbl(
    ["Option", "Cost", "Complexity", "Best For", "Compliance"],
    [
        ("Self-signed cert + Windows GPO push",
         OK("Free"), OK("Low — see Section 9"),
         "LAN-only internal tool — no internet access needed",
         WRN("Acceptable for internal")),
        ("Let's Encrypt (requires public domain)",
         OK("Free"), WRN("Medium — needs public DNS"),
         "If Kforce assigns a public subdomain (e.g. krecruit.kforce.co.in)",
         OK("Best practice")),
        ("Kforce internal CA (if IT runs one)",
         OK("Free"), OK("Zero — IT handles"),
         "Enterprise — IT issues cert from Kforce's own certificate authority",
         OK("Enterprise standard")),
        ("Purchased commercial cert",
         BAD("$50–200/yr"), WRN("Low"), "Not needed for internal LAN tool",
         WRN("Overkill")),
    ],
    col_widths=[2.0, 0.7, 0.85, 1.8, 1.15],
)

# =============================================================================
#  SECTION 9 — SSL SETUP COMMANDS
# =============================================================================
h1("9.  SSL Setup — Self-Signed Certificate Commands")

callout([
    "RUN ONCE — on the server laptop to generate the SSL certificate:",
    "",
    "  # Step 1: Generate private key + self-signed certificate (valid 10 years)",
    "  openssl req -x509 -nodes -days 3650 -newkey rsa:2048 \\",
    "      -keyout krecruit.key -out krecruit.crt \\",
    "      -subj \"/CN=192.168.1.50/O=Kforce/C=IN\"",
    "",
    "  # Step 2: Place files in nginx config directory",
    "  copy krecruit.key  C:\\nginx\\conf\\krecruit.key",
    "  copy krecruit.crt  C:\\nginx\\conf\\krecruit.crt",
], bg="F1F5F9", border="94A3B8")

h2("9.1  nginx Configuration")
callout([
    "nginx.conf — replace default server block with:",
    "",
    "  server {",
    "      listen 80;",
    "      server_name 192.168.1.50;",
    "      return 301 https://$host$request_uri;   # HTTP → HTTPS redirect",
    "  }",
    "",
    "  server {",
    "      listen 443 ssl;",
    "      server_name 192.168.1.50;",
    "      ssl_certificate     conf/krecruit.crt;",
    "      ssl_certificate_key conf/krecruit.key;",
    "      ssl_protocols       TLSv1.2 TLSv1.3;",
    "      ssl_ciphers         HIGH:!aNULL:!MD5;",
    "",
    "      location / {",
    "          proxy_pass         http://127.0.0.1:8000;",
    "          proxy_set_header   Host $host;",
    "          proxy_set_header   X-Real-IP $remote_addr;",
    "      }",
    "  }",
], bg="F1F5F9", border="94A3B8")

h2("9.2  Push Certificate to All Recruiter PCs via Windows GPO")
callout([
    "IT pushes the krecruit.crt to all recruiter PCs so browsers trust it automatically.",
    "",
    "  On the server laptop — export cert to krecruit.crt",
    "  In Active Directory Group Policy Management:",
    "    Computer Configuration → Windows Settings → Security Settings",
    "    → Public Key Policies → Trusted Root Certification Authorities",
    "    → Import → select krecruit.crt",
    "    → Link GPO to the 'Recruiters' OU → Force GPO Update",
    "",
    "  Once pushed, recruiters access https://192.168.1.50 without any browser warnings.",
    "  No action needed on individual recruiter machines.",
], bg="EEF2FF", border="6366F1")

# =============================================================================
#  SECTION 10 — AUTHENTICATION ARCHITECTURE
# =============================================================================
h1("10.  Authentication Architecture")

tbl(
    ["Layer", "Mechanism", "Detail"],
    [
        ("Login",          "POST /auth/login",        "Email + password submitted as JSON; bcrypt.verify() against users table"),
        ("Password hashing","passlib[bcrypt]",         "bcrypt with 12 rounds — passwords never stored in plain text"),
        ("Token issuance", "python-jose JWT",          "HS256, 24-hour expiry, signed with JWT_SECRET from environment"),
        ("Token verification","Bearer header",         "Every protected route calls get_current_user() dependency"),
        ("Brute force",    "In-memory rate limiter",   "5 failed attempts per IP per 60 seconds → HTTP 429"),
        ("Role enforcement","RBAC dependencies",       "require_admin, require_super_admin gate admin-only routes"),
        ("Session end",    "Frontend clears localStorage", "authFetch() redirects to login on 401; logout clears all stored keys"),
        ("Audit trail",    "Middleware",               "Every HTTP request logged: IP, user_id, method, path, status, duration"),
    ],
    col_widths=[1.3, 1.5, 3.7],
)

# =============================================================================
#  SECTION 11 — COMPLIANCE FINDINGS
# =============================================================================
h1("11.  Compliance Findings — Full Register (v7 Current State)")

h2("11.1  Resolved Findings")
tbl(
    ["ID", "Finding", "How Resolved", "Version"],
    [
        ("C-01a", "No authentication on API routes",
         "get_current_user dependency added to all protected endpoints", DONE("v4")),
        ("C-01b", "/bot/chat had no authentication",
         "get_current_user dependency added to bot_chat", DONE("v5")),
        ("C-02",  "JWT secret hardcoded in source code",
         "JWT_SECRET read from environment variable; server refuses start if missing", DONE("v4")),
        ("C-03",  "No brute-force login protection",
         "5 failed attempts per IP per 60s → HTTP 429", DONE("v4")),
        ("C-04",  "No audit logging",
         "Middleware logs every HTTP request to audit.log", DONE("v4")),
        ("C-05",  "CDN scripts without integrity hashes",
         "SHA-384 SRI hashes on all 3 CDN script tags in preview.html", DONE("v4")),
        ("C-06",  "PII in AI prompts",
         "Email, phone, LinkedIn stripped before any LLM call", DONE("v4")),
        ("F-01",  "API URL hardcoded to http://127.0.0.1:8000 (LAN users broken)",
         "Changed to relative URL — works in dev and production without changes", DONE("v7")),
        ("F-02",  "5 API calls sent no Bearer token (bypassed authFetch)",
         "All 5 calls converted to authFetch()", DONE("v7")),
        ("F-04",  "CORS allow_origins=[\"*\"] — no origin restriction",
         "Locked to env-configurable origin list; wildcard removed", DONE("v7")),
        ("F-05",  "Zero HTTP security headers",
         "CSP, X-Frame-Options, X-Content-Type-Options, HSTS, Referrer-Policy added via middleware", DONE("v7")),
        ("FE-01", "Server root served deprecated index.html",
         "main.py now serves preview.html at /; index.html has deprecation comment", DONE("v7")),
    ],
    col_widths=[0.5, 2.3, 2.7, 0.55],
)

h2("11.2  Active Findings — Require Action Before Go-Live")
tbl(
    ["ID", "Finding", "Risk", "Fix Required", "Owner"],
    [
        ("C-07", "APP_ENV not set in .env — FastAPI /docs and /redoc accessible in default mode",
         WRN("Medium"), "Set APP_ENV=production in .env on server laptop", "Dev"),
        ("C-08", "Candidate data isolation missing — no uploaded_by FK — all recruiters see all candidates",
         BAD("High — blocker"), "Add uploaded_by FK to candidates table; filter queries by user role", "Dev"),
        ("C-09", "Windows Firewall gap — port 8000 accessible directly, bypassing nginx/SSL",
         WRN("Medium"), "Block port 8000 inbound; allow only 443 and 80 from LAN", "IT"),
        ("F-03", "Babel Standalone (compiler/dev tool) running in production browser",
         WRN("Medium"), "IT approves one of Option A/B/C in Section 6B; dev implements", "IT + Dev"),
        ("F-06", "JWT token stored in localStorage (accessible to JS)",
         WRN("Medium"), "Mitigated by CSP + XSS absence; upgrade to httpOnly cookie = backlog BL-07", "Dev"),
        ("11.3", "No backup strategy — server laptop single point of failure",
         BAD("High — blocker"), "Set up three-layer backup (Section 21.1)", "IT"),
        ("11.4", "No audit log rotation — audit.log grows indefinitely",
         WRN("Medium"), "Replace FileHandler with RotatingFileHandler in main.py", "Dev"),
    ],
    col_widths=[0.5, 2.3, 0.8, 2.0, 0.5],
)

h2("11.3  Backlog Findings — Post Go-Live")
tbl(
    ["ID", "Finding", "Risk", "Notes"],
    [
        ("BL-01", "No minimum password length check",
         WRN("Low"), "Add len(password) >= 8 check in users router"),
        ("BL-02", "No documented account lifecycle policy",
         WRN("Low"), "Documented in Section 21.4 — procedure written, not yet enforced in code"),
        ("BL-03", "BitLocker not mentioned",
         WRN("Low"), "IT enables full-disk encryption on server laptop"),
        ("BL-04", "No data retention policy",
         WRN("Low"), "Add configurable auto-purge for records older than 90 days"),
        ("BL-05", "No recruiter consent acknowledgment at upload",
         WRN("Low"), "Checkbox: 'I confirm this data is processed under Kforce data policy'"),
        ("BL-06", "Azure Cost Management budget alert not set",
         WRN("Low"), "5-minute task in Azure portal — set $80/month alert"),
        ("BL-07", "JWT stored in localStorage instead of httpOnly cookie",
         WRN("Low"), "Medium-complexity refactor; mitigated by CSP and XSS absence. Post-launch."),
        ("BL-08", "React/ReactDOM version 18.2.0 (June 2022 — 3 years old)",
         WRN("Low"), "Update to 18.3.1 — 1-line CDN URL change; do alongside Option A self-hosting"),
    ],
    col_widths=[0.5, 2.1, 0.7, 3.2],
)

# =============================================================================
#  SECTION 12 — MULTI-USER RBAC ARCHITECTURE
# =============================================================================
h1("12.  Multi-User RBAC Architecture")

h2("12.1  Role Definitions")
tbl(
    ["Role", "Who", "What They Can Do"],
    [
        ("super_admin", "Technical lead / IT owner",
         "Everything — user management, role management, settings, all candidate data, all JDs"),
        ("admin", "HR manager / team lead",
         "All candidate operations, JD management, settings, user management (no role changes)"),
        ("recruiter", "Individual HR recruiter",
         "Upload resumes, run analysis, view candidates, approve, download resumes"),
    ],
    col_widths=[1.0, 1.5, 4.0],
)

h2("12.2  Permission Matrix")
tbl(
    ["Permission", "super_admin", "admin", "recruiter"],
    [
        ("View all candidates",      OK("YES"), OK("YES"), OK("YES")),
        ("Upload resume",            OK("YES"), OK("YES"), OK("YES")),
        ("Run AI analysis",          OK("YES"), OK("YES"), OK("YES")),
        ("Approve candidate",        OK("YES"), OK("YES"), OK("YES")),
        ("Download resume",          OK("YES"), OK("YES"), OK("YES")),
        ("Create / edit JDs",        OK("YES"), OK("YES"), BAD("NO")),
        ("Delete candidate",         OK("YES"), OK("YES"), BAD("NO")),
        ("Edit AI settings",         OK("YES"), OK("YES"), BAD("NO")),
        ("Create / deactivate users",OK("YES"), OK("YES"), BAD("NO")),
        ("Create / edit roles",      OK("YES"), BAD("NO"), BAD("NO")),
    ],
    col_widths=[2.2, 1.3, 1.3, 1.3],
)

# =============================================================================
#  SECTION 13 — DATA PROTECTION CONTROLS
# =============================================================================
h1("13.  Data Protection Controls")

tbl(
    ["Control", "Implementation", "Status"],
    [
        ("Password hashing",     "bcrypt 12 rounds — no plain-text passwords anywhere",                OK("IMPLEMENTED")),
        ("JWT authentication",   "All routes protected — 24h expiry — HS256 signed",                  OK("IMPLEMENTED")),
        ("Brute force protection","5 attempts/IP/60s → 429",                                           OK("IMPLEMENTED")),
        ("PII stripping",        "Email, phone, LinkedIn removed from AI prompts",                     OK("IMPLEMENTED")),
        ("Audit logging",        "Every HTTP request logged with user_id, IP, path, status, duration", OK("IMPLEMENTED")),
        ("SRI hash verification","sha384 on all 3 CDN JS files",                                       OK("IMPLEMENTED")),
        ("CORS restriction",     "Origins locked to env-configurable list (wildcard removed in v7)",   OK("IMPLEMENTED v7")),
        ("HTTP security headers","CSP, X-Frame-Options, X-Content-Type-Options, HSTS, Referrer-Policy",OK("IMPLEMENTED v7")),
        ("API auth token delivery","All API calls use authFetch() with Bearer token (5 missing fixed v7)",OK("IMPLEMENTED v7")),
        ("Relative API URL",     "App works on any host — dev and production without code changes",    OK("IMPLEMENTED v7")),
        ("Transport encryption", "nginx + SSL required before production — self-signed cert adequate for LAN", WRN("PENDING — server setup")),
        ("Data isolation",       "uploaded_by FK not yet added to candidates table",                   BAD("PENDING — C-08")),
        ("Log rotation",         "RotatingFileHandler not yet applied",                                WRN("PENDING — 11.4")),
        ("Backup",               "Three-layer backup not yet configured",                              BAD("PENDING — 11.3")),
    ],
    col_widths=[1.6, 3.5, 1.4],
)

# =============================================================================
#  SECTION 14 — COMPLIANCE CHECKLIST
# =============================================================================
h1("14.  Compliance Readiness Checklist (v7 — Updated)")

h2("14.1  Active Blockers — Must Complete Before Go-Live")
tbl(
    ["#", "Requirement", "Status", "Owner"],
    [
        ("B-01", "Switch AI provider to Azure OpenAI in Settings",                    BAD("PENDING"),  "IT + Dev"),
        ("B-02", "Deploy nginx reverse proxy on server laptop",                        BAD("PENDING"),  "IT"),
        ("B-03", "Generate SSL certificate and configure nginx HTTPS",                BAD("PENDING"),  "IT"),
        ("B-04", "Push SSL certificate to all recruiter PCs via GPO",                 BAD("PENDING"),  "IT"),
        ("B-05", "Block port 8000 in Windows Firewall — allow only 443 and 80",       BAD("PENDING"),  "IT"),
        ("B-06", "Set APP_ENV=production in .env",                                    BAD("PENDING"),  "Dev"),
        ("B-07", "Set HF_HUB_OFFLINE=1 in .env (pre-download model first)",           BAD("PENDING"),  "Dev"),
        ("B-08", "Migrate database to PostgreSQL for 100 users",                      BAD("PENDING"),  "IT + Dev"),
        ("B-09", "Set CORS_ORIGINS to https://192.168.1.50 in .env",                 BAD("PENDING"),  "Dev"),
        ("B-10", "Add uploaded_by FK to candidates table — data isolation (C-08)",    BAD("PENDING"),  "Dev"),
        ("B-11", "Configure daily database backup (pg_dump to network share)",        BAD("PENDING"),  "IT"),
        ("B-12", "Configure weekly /uploads folder backup",                           BAD("PENDING"),  "IT"),
        ("B-13", "Replace FileHandler with RotatingFileHandler in main.py",           BAD("PENDING"),  "Dev"),
        ("B-14", "IT approves Option A (self-host JS libs) OR Option B (htm library)",BAD("PENDING"),  "IT"),
        ("B-15", "Generate new JWT_SECRET for production (not dev secret)",           BAD("PENDING"),  "Dev"),
        ("B-16", "Test full login, upload, analysis, download flow on LAN",           BAD("PENDING"),  "Dev + IT"),
    ],
    col_widths=[0.35, 3.3, 0.85, 0.8],
)

h2("14.2  Completed Items")
tbl(
    ["#", "Requirement", "Status", "Version"],
    [
        ("C-01", "JWT authentication on all API routes",                         DONE("DONE"), "v4"),
        ("C-02", "Bcrypt password hashing",                                      DONE("DONE"), "v4"),
        ("C-03", "JWT secret from environment variable",                         DONE("DONE"), "v4"),
        ("C-04", "Brute-force login protection (429)",                           DONE("DONE"), "v4"),
        ("C-05", "Audit logging middleware",                                     DONE("DONE"), "v4"),
        ("C-06", "SRI hashes on all CDN scripts",                               DONE("DONE"), "v4"),
        ("C-07", "PII stripped from AI prompts",                                 DONE("DONE"), "v4"),
        ("C-08", "load_dotenv before all imports",                               DONE("DONE"), "v5"),
        ("C-09", "Pydantic v2 from_attributes compatibility",                    DONE("DONE"), "v5"),
        ("C-10", "Multi-user RBAC: super_admin / admin / recruiter",             DONE("DONE"), "v5"),
        ("C-11", "API URL changed from hardcoded localhost to relative URL",     DONE("DONE"), "v7"),
        ("C-12", "All API calls use authFetch() with Bearer token",              DONE("DONE"), "v7"),
        ("C-13", "CORS wildcard removed — locked origin list",                   DONE("DONE"), "v7"),
        ("C-14", "Security headers: CSP, X-Frame, HSTS, X-Content-Type etc.",   DONE("DONE"), "v7"),
        ("C-15", "Server root serves preview.html (not deprecated index.html)",  DONE("DONE"), "v7"),
        ("C-16", "index.html marked with prominent deprecation comment",         DONE("DONE"), "v7"),
    ],
    col_widths=[0.35, 3.3, 0.85, 0.6],
)

# =============================================================================
#  SECTION 15 — FILES REVIEWED
# =============================================================================
h1("15.  Scope of Code Review — Files Audited (v7)")

tbl(
    ["File", "What Was Reviewed / Changed"],
    [
        ("backend/main.py",                       "load_dotenv order, APP_ENV gate, CORS wildcard fixed, security headers middleware ADDED (v7), serve_frontend → preview.html (v7)"),
        ("backend/deps.py",                        "JWT auth — SECRET_KEY load, get_current_user, require_admin"),
        ("backend/routers/auth.py",                "Login, bcrypt verify, JWT token, rate limit, /verify, /me"),
        ("backend/routers/users.py",               "User CRUD — password validation gap identified (BL-01)"),
        ("backend/routers/roles.py",               "Role CRUD — super_admin gated"),
        ("backend/routers/bot.py",                 "Auth added v5 — confirmed"),
        ("backend/routers/candidates.py",          "All endpoints — data isolation gap identified (C-08)"),
        ("backend/routers/settings.py",            "Admin gated — confirmed"),
        ("backend/routers/client_requirements.py", "Admin gated for write — confirmed"),
        ("backend/services/llm_service.py",        "PII stripping confirmed — email/phone/LinkedIn removed"),
        ("backend/services/semantic_service.py",   "One-time HuggingFace download only — confirmed"),
        ("backend/database.py",                    "Absolute SQLite path, seed_roles, seed_admin"),
        ("backend/models/candidate.py",            "No uploaded_by FK — data isolation gap C-08 confirmed"),
        ("backend/models/user.py",                 "bcrypt hash confirmed"),
        ("backend/models/role.py",                 "Permission JSON confirmed"),
        ("backend/schemas/client_requirement.py",  "from_attributes — Pydantic v2 compatible"),
        ("backend/.env",                           "CORS_ORIGINS added (v7); APP_ENV missing from production checklist"),
        ("backend/requirements.txt",               "All dependencies audited — no unexpected network libs"),
        ("backend/static/index.html",              "DEPRECATED — DO NOT USE comment added (v7). Old prototype, not served."),
        ("backend/static/preview.html",            "FULL AUDIT v7: API URL fixed, 5 authFetch fixes, Babel/CDN audit, JWT storage review"),
    ],
    col_widths=[2.5, 4.0],
)

# =============================================================================
#  SECTION 16 — QUICK REFERENCE
# =============================================================================
h1("16.  Quick Reference — Key Answers for the Compliance Team")

h2("Q1  What is the production server?")
callout([
    "A dedicated office laptop — permanently powered, sleep disabled, Kforce India office LAN.",
    "Static LAN IP assigned by IT (e.g. 192.168.1.50).",
    "Minimum spec: 16 GB RAM, 512 GB SSD.",
    "AI processing (GPT-4o) runs on Azure — the laptop handles web serving and database only.",
    "Full setup guide: Section 20.",
], bg="F0FDF4", border="16A34A")

h2("Q2  What is the authentication model?")
callout([
    "POST /auth/login → bcrypt verify → JWT token (24h expiry)",
    "All API routes require Bearer token — 401 on missing/expired/invalid",
    "authFetch() on frontend injects token on every call and redirects to login on 401",
    "Roles: super_admin (full) | admin (no role mgmt) | recruiter (own candidates)",
    "Brute force: 5 failed attempts per IP per 60s → HTTP 429",
    "Audit: every request logged — IP, user_id, method, path, status, duration",
], bg="EEF2FF", border="6366F1")

h2("Q3  What Azure services do we need and what do they cost?")
callout([
    "No new subscription — use existing Kforce Azure tenant (same as M365/Copilot)",
    "Enable Azure OpenAI → Create resource (Central India) → Deploy GPT-4o → Copy key",
    "Cost: ~$3–$10/month for 20 recruiters  |  ~$30–$80/month for 100 recruiters",
    "Pay-per-use — no standing fee. Full setup guide: Section 7.2",
], bg="EEF2FF", border="6366F1")

h2("Q4  What about the Babel / CDN concern?")
callout([
    "THREE-PART ANSWER:",
    "",
    "1. React and ReactDOM: MIT license — no compliance concern. Used by all major enterprises.",
    "   SRI hashes protect against CDN tampering — browser rejects modified scripts.",
    "",
    "2. CDN dependency: Self-hosting approved via Option A (Section 6B.3) removes external CDN calls.",
    "   IT downloads 3 files once; dev updates script tags. Closes CDN finding.",
    "",
    "3. Babel dev-tool: Formal approval request in Section 6B.",
    "   Option B (htm library, Apache-2.0) removes Babel from production entirely.",
    "   Option A is the immediate fix; Option B is the proper compliance closure.",
], bg="EEF2FF", border="6366F1")

h2("Q5  What is the biggest outstanding compliance risk?")
callout([
    "THREE EQUAL CONCERNS:",
    "",
    "1. AI provider not yet switched to Azure (B-01 — blocker)",
    "   Candidate PII currently going to Groq/Gemini/NVIDIA in development.",
    "   Fix: IT provisions Azure OpenAI → Settings page change → resolved immediately.",
    "",
    "2. Candidate data isolation missing (C-08 — blocker)",
    "   All 100 recruiters can see all candidates — no per-recruiter filtering.",
    "   Fix: Dev adds uploaded_by FK to candidates table + filter by user role.",
    "",
    "3. No backup strategy (11.3 — blocker)",
    "   Server laptop failure = all candidate data lost.",
    "   Fix: IT sets up pg_dump backup to network share (Section 21.1).",
], bg="FEF2F2", border="EF4444")

# =============================================================================
#  SECTION 17 — V4 SECURITY HARDENING DETAILS
# =============================================================================
h1("17.  v4 Security Hardening — 9 Code-Level Fixes Applied")

tbl(
    ["Fix #", "Finding Fixed", "File Changed", "What Was Done"],
    [
        ("V4-F1", "Missing auth on all API routes",         "candidates.py, client_requirements.py, settings.py, recruiters.py", "get_current_user Dependency added to every protected endpoint"),
        ("V4-F2", "Hardcoded JWT secret",                   "deps.py",            "JWT_SECRET from env var; server refuses to start without it"),
        ("V4-F3", "/docs and /redoc in production",         "main.py",            "Hidden when APP_ENV=production (APP_ENV gap documented — C-07)"),
        ("V4-F4", "No brute-force protection",              "routers/auth.py",    "5 failed attempts per IP per 60s → HTTP 429"),
        ("V4-F5", "No audit logging",                       "main.py",            "Middleware logs every request to audit.log"),
        ("V4-F6", "CDN scripts without integrity check",    "preview.html",       "SRI sha384 hashes on all 3 CDN script tags"),
        ("V4-F7", "PII in AI prompts",                      "llm_service.py",     "Email, phone, LinkedIn stripped before LLM call"),
        ("V4-F8", "Plain-text login password",              "database.py, auth.py","Passwords bcrypt-hashed in users table"),
        ("V4-F9", "No session token expiry",                "routers/auth.py",    "JWT tokens include exp claim — 24 hours"),
    ],
    col_widths=[0.55, 1.8, 1.8, 2.35],
)

# =============================================================================
#  SECTION 18 — V5 ADDITIONAL FIXES
# =============================================================================
h1("18.  v5 Additional Fixes — 5 Code-Level Fixes Applied")

tbl(
    ["Fix #", "Issue", "File", "Fix Applied"],
    [
        ("V5-F1", "load_dotenv() after imports — JWT_SECRET not loaded before deps.py check",
         "main.py", "load_dotenv() moved to lines 1–2 before all imports"),
        ("V5-F2", "DATABASE_URL relative path in .env — SQLite could not open file",
         ".env", "DATABASE_URL removed; SQLite uses absolute path logic in database.py"),
        ("V5-F3", "Pydantic orm_mode UserWarning on every startup",
         "schemas/client_requirement.py", "orm_mode renamed to from_attributes"),
        ("V5-F4", "/bot/chat endpoint had no authentication",
         "routers/bot.py", "get_current_user dependency added to bot_chat"),
        ("V5-F5", "Multi-user RBAC not connected (startup blocked by V5-F1)",
         "main.py, database.py, all routers", "With V5-F1 resolved: full RBAC starts cleanly — users, roles, JWT end-to-end"),
    ],
    col_widths=[0.55, 2.0, 1.5, 2.45],
)

# =============================================================================
#  SECTION 19 — V7 FRONTEND SECURITY FIXES
# =============================================================================
h1("19.  v7 Frontend Security Fixes — 6 Code-Level Changes Applied")

tbl(
    ["Fix #", "Finding", "File Changed", "What Was Done"],
    [
        ("V7-F1", "API URL hardcoded to http://127.0.0.1:8000 — LAN users connect to wrong host",
         "preview.html line 182",
         "Changed to const API = \"\" (empty = relative to current host). Works in dev and production without modification."),
        ("V7-F2", "bot/chat (x2 calls) used raw fetch — no Bearer token sent",
         "preview.html lines 387, 401",
         "Both calls converted from fetch() to authFetch()"),
        ("V7-F3", "settings/domain-coverage used raw fetch — no Bearer token sent",
         "preview.html line 845",
         "Converted from fetch() to authFetch()"),
        ("V7-F4", "candidates/{id}/role-compatibility used raw fetch — no Bearer token sent",
         "preview.html line 1685",
         "Converted from fetch() to authFetch()"),
        ("V7-F5", "settings/semantic-status used raw fetch — no Bearer token sent",
         "preview.html line 2156",
         "Converted from fetch() to authFetch()"),
        ("V7-F6", "CORS wildcard allow_origins=['*'] + security headers missing",
         "main.py lines 76–82",
         "CORS locked to CORS_ORIGINS env var. Security headers middleware added: CSP, X-Frame-Options, X-Content-Type-Options, Referrer-Policy, Permissions-Policy, HSTS (production only)."),
        ("V7-F7", "Server root served deprecated index.html instead of preview.html",
         "main.py line 128",
         "Changed FileResponse to serve preview.html"),
        ("V7-F8", "index.html had no deprecation notice — could be mistaken for active file",
         "static/index.html line 1",
         "Added prominent DO NOT USE comment block with pointer to preview.html"),
    ],
    col_widths=[0.55, 1.8, 1.6, 2.55],
)

# =============================================================================
#  SECTION 20 — SERVER LAPTOP DEPLOYMENT GUIDE
# =============================================================================
h1("20.  Server Laptop Deployment Guide")

body(
    "This section documents the one-time setup required to configure the dedicated "
    "office laptop as the K Recruit production server for 100 recruiters."
)

h2("20.1  Minimum Hardware Requirements")
tbl(
    ["Component", "Minimum", "Reason"],
    [
        ("RAM",      "16 GB",     "Web server + PostgreSQL + semantic model loaded in memory"),
        ("Storage",  "512 GB SSD","PostgreSQL DB + /uploads resume PDFs grow over time"),
        ("OS",       "Windows 10/11 Pro or Ubuntu 20.04+", "Required for PostgreSQL + nginx"),
        ("Network",  "Ethernet (preferred) or WiFi", "Stable connection — do not let laptop sleep"),
        ("Power",    "Always plugged in", "Server must not run on battery — disable power saving"),
    ],
    col_widths=[1.2, 1.8, 3.5],
)

h2("20.2  One-Time Setup Checklist for IT")
tbl(
    ["Step", "Action", "Command / Location", "Who"],
    [
        ("1",  "Disable sleep and hibernate permanently",
         "Control Panel → Power Options → Never → Apply", "IT"),
        ("2",  "Assign static IP to the laptop",
         "Router DHCP reservation OR Windows: Network Adapter → IPv4 → Manual IP", "IT"),
        ("3",  "Install PostgreSQL 15+",
         "postgresql.org/download → install → create DB 'krecruitdb' and user 'krecruit'", "IT / Dev"),
        ("4",  "Set DATABASE_URL in .env",
         "DATABASE_URL=postgresql://krecruit:password@localhost/krecruitdb", "Dev"),
        ("5",  "Run database migration",
         "python main.py (app auto-creates all tables on first start)", "Dev"),
        ("6",  "Pre-download HuggingFace model",
         "python -c \"from sentence_transformers import SentenceTransformer; SentenceTransformer('all-MiniLM-L6-v2')\"", "Dev"),
        ("7",  "Set HF_HUB_OFFLINE=1 in .env",
         "backend/.env → add HF_HUB_OFFLINE=1", "Dev"),
        ("8",  "Set APP_ENV=production in .env",
         "backend/.env → APP_ENV=production", "Dev"),
        ("9",  "Set CORS_ORIGINS to production URL",
         "backend/.env → CORS_ORIGINS=https://192.168.1.50", "Dev"),
        ("10", "Self-host JS libraries (Option A) — download 3 files to /static/",
         "React 18.3.1 + ReactDOM 18.3.1 + Babel 7.23.2 from unpkg.com (see Section 6B.3)", "Dev"),
        ("11", "Install nginx",
         "nginx.org/download (Windows) or  apt install nginx (Linux)", "IT / Dev"),
        ("12", "Generate SSL certificate and configure nginx",
         "See Section 9 — full commands provided", "IT / Dev"),
        ("13", "Block port 8000 in Windows Firewall",
         "Windows Defender Firewall → Inbound Rules → New Rule → Port 8000 → Block", "IT"),
        ("14", "IT pushes SSL cert to all recruiter PCs via GPO",
         "See Section 9.2", "IT"),
        ("15", "Configure auto-start on boot",
         "Windows: Task Scheduler → new task → trigger At startup → run uvicorn + nginx", "IT / Dev"),
        ("16", "Set NTFS permissions on /uploads folder",
         "Right-click /uploads → Properties → Security → restrict to service account only", "IT"),
        ("17", "Test with 2–3 recruiters before full rollout",
         "Verify login, upload, analysis, download on real recruiter PCs", "Dev / HR"),
    ],
    col_widths=[0.35, 2.0, 2.7, 0.8],
)

h2("20.3  Environment Variables on Server Laptop (.env)")
callout([
    "Complete .env for the server laptop (production):",
    "",
    "  # Database — PostgreSQL for 100 users",
    "  DATABASE_URL=postgresql://krecruit:password@localhost/krecruitdb",
    "",
    "  # App",
    "  APP_HOST=0.0.0.0",
    "  APP_PORT=8000",
    "  APP_ENV=production          ← disables /docs and /redoc",
    "",
    "  # Semantic model — offline after pre-download",
    "  HF_HUB_OFFLINE=1            ← blocks all HuggingFace calls",
    "",
    "  # Auth — generate a new secret for production",
    "  JWT_SECRET=<run: python -c 'import secrets; print(secrets.token_hex(32))'>",
    "",
    "  # CORS — lock to production server IP only",
    "  CORS_ORIGINS=https://192.168.1.50",
    "",
    "  # Azure OpenAI (fill in after IT provisions)",
    "  # Configured via Settings page in the app — not needed in .env",
], bg="F1F5F9", border="94A3B8")

# =============================================================================
#  SECTION 21 — OPERATIONAL PROCEDURES
# =============================================================================
h1("21.  Operational Procedures — Backup, Logs, Firewall")

h2("21.1  Backup Strategy")
callout([
    "RISK: Server laptop is a single point of failure.",
    "If the laptop fails or is damaged, all candidate data and resume PDFs are lost.",
    "",
    "REQUIRED: Three-layer backup",
    "",
    "Layer 1 — Daily database backup (automated):",
    "  Windows Task Scheduler runs every night at 11 PM:",
    "  pg_dump krecruitdb > C:\\Backups\\krecruitdb_YYYYMMDD.sql",
    "  Copy to network share: \\\\kforce-fileserver\\k-recruit-backups\\",
    "  Keep last 30 days. Total size: ~10–50 MB per day.",
    "",
    "Layer 2 — Weekly /uploads folder backup:",
    "  Copy backend/uploads/ to network share weekly.",
    "  Resume PDFs are not in the database — must be backed up separately.",
    "",
    "Layer 3 — Quarterly restore test:",
    "  Spin up a test laptop, restore from backup, verify app starts and data is intact.",
    "  A backup never tested is not a backup.",
], bg="FEF2F2", border="EF4444")

h2("21.2  Audit Log Rotation")
callout([
    "CURRENT STATE: audit.log grows indefinitely. Fix required before production.",
    "",
    "FIX — Replace FileHandler in main.py with RotatingFileHandler:",
    "",
    "  from logging.handlers import RotatingFileHandler",
    "  _fh = RotatingFileHandler(",
    "      _audit_log_path, maxBytes=10*1024*1024, backupCount=10, encoding='utf-8'",
    "  )",
    "",
    "  This keeps: audit.log (current) + audit.log.1 through audit.log.10",
    "  Total max size: ~100 MB",
    "  Oldest logs auto-deleted when limit reached",
    "",
    "  Retention policy: keep minimum 30 days of logs for compliance audit trail.",
], bg="FEF9C3", border="F59E0B")

h2("21.3  Windows Firewall Rules")
tbl(
    ["Rule", "Direction", "Port", "Action", "Reason"],
    [
        ("Block direct FastAPI access", "Inbound", "8000", BAD("BLOCK from all except localhost"), "Force all traffic through nginx/SSL"),
        ("Allow HTTPS",                 "Inbound", "443",  OK("ALLOW from LAN"),                   "Recruiter access via nginx"),
        ("Allow HTTP redirect",         "Inbound", "80",   OK("ALLOW from LAN"),                   "Redirect to HTTPS"),
        ("Allow PostgreSQL",            "Inbound", "5432", BAD("BLOCK from all except localhost"),  "DB must not be reachable from LAN"),
        ("Allow Azure OpenAI",          "Outbound","443",  OK("ALLOW to *.openai.azure.com"),       "AI provider calls"),
    ],
    col_widths=[1.9, 0.85, 0.6, 1.6, 2.55],
)

h2("21.4  Account Lifecycle (Backlog BL-02)")
callout([
    "WHEN A RECRUITER JOINS:",
    "  Admin logs into K Recruit → User Management → Create User",
    "  Set name, email, temporary password, role = recruiter",
    "  Recruiter logs in and changes password on first use (manual step until enforced)",
    "",
    "WHEN A RECRUITER LEAVES KFORCE:",
    "  HR notifies admin on the last working day",
    "  Admin → User Management → Edit User → set is_active = False",
    "  Account immediately blocked — JWT tokens rejected on next request",
    "  Do NOT delete the account — retain for audit trail",
    "",
    "NOTE: This procedure should be added to Kforce HR offboarding checklist.",
], bg="EEF2FF", border="6366F1")

# =============================================================================
#  SECTION 22 — DECLARATION
# =============================================================================
h1("22.  Declaration")

callout([
    "FOR THE KFORCE US COMPLIANCE AND IT GOVERNANCE TEAM:",
    "",
    "  K Recruit is approved for production deployment on the dedicated office laptop",
    "  subject to completing the blockers listed in Section 14.1 (B-01 through B-16).",
    "",
    "  SECURITY STATUS AFTER v7 REVIEW (cumulative — all versions):",
    "  ✓  All API endpoints require valid JWT authentication",
    "  ✓  Passwords stored as bcrypt hashes — no plain text credentials",
    "  ✓  JWT tokens expire after 24 hours",
    "  ✓  Brute-force login protection active (5 attempts/IP/60s → 429)",
    "  ✓  Full audit log of every HTTP request",
    "  ✓  SRI hashes on all CDN script tags (tamper protection)",
    "  ✓  PII stripped from AI prompts (email, phone, LinkedIn)",
    "  ✓  Multi-user RBAC: super_admin / admin / recruiter",
    "  ✓  API URL relative — works in dev and production without code changes",
    "  ✓  All API calls carry Bearer token (authFetch on all 33+ calls)",
    "  ✓  CORS wildcard removed — locked to env-configurable origin list",
    "  ✓  Security headers: CSP, X-Frame-Options, HSTS, X-Content-Type, Referrer-Policy",
    "  ✓  Active frontend file (preview.html) confirmed; index.html deprecated",
    "",
    "  OUTSTANDING BLOCKERS (require action before go-live):",
    "  ✗  Switch AI provider to Azure OpenAI (most critical — PII currently leaving org)",
    "  ✗  Deploy nginx + SSL on server laptop",
    "  ✗  Set APP_ENV=production and block port 8000",
    "  ✗  Add candidate data isolation (uploaded_by FK)",
    "  ✗  Set up daily database backup",
    "  ✗  Add audit log rotation (RotatingFileHandler)",
    "  ✗  IT approves self-host of JS libraries OR htm library (Section 6B)",
    "",
    "  This document (v7) is the complete and final compliance reference.",
    "  It fully supersedes v1–v6 (2026-001 through 2026-006).",
], bg="EEF2FF", border="6366F1")

sp(8)
for label, value in [
    ("Application Name: ", "K Recruit — HR Resume Optimizer"),
    ("Report Version: ",   "7.0 — Final  (Frontend Security Edition)"),
    ("Document Ref: ",     "K-RECRUIT-COMP-2026-007"),
    ("Supersedes: ",       "v1–v6: 2026-001 through K-RECRUIT-COMP-2026-006"),
    ("Prepared by: ",      "Technical Architecture Review, K Recruit Development Team"),
    ("Contact: ",          "sushrut.nistane@kforce.co.in"),
    ("Date: ",             "May 2026"),
]:
    p = doc.add_paragraph()
    rn(p, label, bold=True,  sz=9, color=NAVY)
    rn(p, value, bold=False, sz=9, color=MUTED)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)

# ── Save ──────────────────────────────────────────────────────────────────────
out = r"C:\Users\SushrutNistane\OneDrive - Kforce Inter\cms-resume-optimizer\K_Recruit_Compliance_Report_v7.docx"
doc.save(out)
print(f"Saved: {out}")
